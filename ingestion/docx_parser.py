"""
P8-A: DOCX Document Parser.

Parses .docx files into the common parsed-document schema defined in
specs/cross_document_bridge.py. Extracts paragraphs, tables, headings,
and document metadata.

DOCX ONLY -- no PDF. Keyword heuristics only -- no LLM.
"""

from __future__ import annotations

import os
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from specs.cross_document_bridge import DocumentParser, DocumentType


class DOCXDocumentParser(DocumentParser):
    """Parse a DOCX file into the common parsed-document schema.

    Extracts:
    - All paragraphs with heading levels
    - All tables with row/col data
    - Document metadata (title, author, created date)
    - Section structure (heading hierarchy)
    """

    def parse(self, document_path: str) -> Dict[str, Any]:
        """Read DOCX and extract: paragraphs, tables, headings, metadata."""
        if not os.path.isfile(document_path):
            raise FileNotFoundError(f"DOCX not found: {document_path}")

        doc = DocxDocument(document_path)

        # Extract metadata
        metadata = self._extract_metadata(doc)

        # Extract paragraphs with heading levels
        paragraphs = self._extract_paragraphs(doc)

        # Extract tables
        tables = self._extract_tables(doc)

        # Build section hierarchy from headings
        sections = self._build_section_hierarchy(paragraphs)

        # Combine paragraph text per "page" (DOCX doesn't have pages,
        # so we treat the entire document as page 1)
        full_text = "\n".join(p["text"] for p in paragraphs if p["text"].strip())

        return {
            "document_path": document_path,
            "document_type": DocumentType.COMPARABILITY_PROTOCOL,
            "pages": [
                {
                    "page_number": 1,
                    "text": full_text,
                    "tables": tables,
                }
            ],
            "paragraphs": paragraphs,
            "sections": sections,
            "metadata": metadata,
        }

    def supported_formats(self) -> List[str]:
        return [".docx"]

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _extract_metadata(self, doc: DocxDocument) -> Dict[str, Any]:
        """Extract document core properties."""
        props = doc.core_properties
        created = props.created
        modified = props.modified

        return {
            "source_format": "docx",
            "title": props.title or "",
            "author": props.author or "",
            "created_date": created.isoformat() if isinstance(created, datetime) else str(created or ""),
            "modified_date": modified.isoformat() if isinstance(modified, datetime) else str(modified or ""),
            "subject": props.subject or "",
            "keywords": props.keywords or "",
        }

    def _extract_paragraphs(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract all paragraphs with heading level info."""
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""
            heading_level = self._get_heading_level(style_name)

            paragraphs.append({
                "index": i,
                "text": para.text,
                "style": style_name,
                "heading_level": heading_level,
                "is_heading": heading_level is not None,
                "runs": [
                    {
                        "text": run.text,
                        "bold": run.bold or False,
                        "italic": run.italic or False,
                    }
                    for run in para.runs
                ],
            })
        return paragraphs

    def _get_heading_level(self, style_name: str) -> Optional[int]:
        """Return heading level (1-9) or None if not a heading."""
        if not style_name:
            return None
        lower = style_name.lower()
        if lower == "title":
            return 0
        if lower.startswith("heading"):
            # "Heading 1", "Heading 2", etc.
            parts = lower.split()
            if len(parts) == 2:
                try:
                    return int(parts[1])
                except ValueError:
                    pass
        return None

    def _extract_tables(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """Extract all tables with headers and row data."""
        tables = []
        for t_idx, table in enumerate(doc.tables):
            rows_data = []
            headers = []

            for r_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if r_idx == 0:
                    headers = cells
                else:
                    # Build a dict mapping header->value for each data row
                    row_dict = {}
                    for c_idx, cell_text in enumerate(cells):
                        if c_idx < len(headers):
                            row_dict[headers[c_idx]] = cell_text
                        else:
                            row_dict[f"col_{c_idx}"] = cell_text
                    rows_data.append(row_dict)

            tables.append({
                "id": f"docx_table_{t_idx + 1}",
                "table_index": t_idx,
                "headers": headers,
                "rows": rows_data,
                "n_rows": len(rows_data),
                "n_cols": len(headers),
            })
        return tables

    def _build_section_hierarchy(
        self, paragraphs: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Build section hierarchy from heading paragraphs."""
        sections = []
        current_section: Optional[Dict[str, Any]] = None

        for para in paragraphs:
            if para["is_heading"]:
                section = {
                    "title": para["text"],
                    "level": para["heading_level"],
                    "paragraph_index": para["index"],
                    "children": [],
                }
                if current_section is None or para["heading_level"] <= (current_section.get("level") or 0):
                    sections.append(section)
                    current_section = section
                else:
                    current_section["children"].append(section)

        return sections
