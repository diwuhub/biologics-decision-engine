"""
ICH Q5E-Compliant Comparability Report Generator (DOCX).

Converts a ComparabilityReport into a professional Word document following
the ICH Q5E structure for comparability assessments of biotechnological
products subject to manufacturing process changes.

Usage:
    from reports.comparability_report import generate_comparability_report
    path = generate_comparability_report(report, "output/report.docx")
"""

from __future__ import annotations

import datetime
import os
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from pipelines.schemas import ComparabilityReport, AttributeResult


# =========================================================================
# Color palette
# =========================================================================

_GREEN = RGBColor(0x27, 0xAE, 0x60)
_YELLOW = RGBColor(0xF3, 0x9C, 0x12)
_RED = RGBColor(0xE7, 0x4C, 0x3C)
_DARK = RGBColor(0x2C, 0x3E, 0x50)
_GRAY = RGBColor(0x7F, 0x8C, 0x8D)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

_BG_GREEN = "D5F5E3"
_BG_YELLOW = "FEF9E7"
_BG_RED = "FADBD8"
_BG_HEADER = "2C3E50"
_BG_LIGHT_GRAY = "F2F3F4"


# =========================================================================
# Concern-to-color mapping
# =========================================================================

def _concern_color(concern: str) -> str:
    """Return background hex color for a concern level."""
    mapping = {
        "none": _BG_GREEN,
        "minor": _BG_GREEN,
        "major": _BG_YELLOW,
        "critical": _BG_RED,
    }
    return mapping.get(concern, _BG_LIGHT_GRAY)


def _concern_action_color(action_level: str) -> str:
    """Return background hex color for an action level."""
    mapping = {
        "PROCEED": _BG_GREEN,
        "SUPPLEMENT": _BG_YELLOW,
        "MONITOR": _BG_YELLOW,
        "INVESTIGATE": _BG_RED,
        "DEFER": _BG_RED,
    }
    return mapping.get(action_level, _BG_LIGHT_GRAY)


def _verdict_color(verdict: str) -> RGBColor:
    """Return text color for a verdict."""
    if verdict == "Comparable":
        return _GREEN
    elif verdict == "Not Comparable":
        return _RED
    return _YELLOW


# =========================================================================
# Table helpers
# =========================================================================

def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 9,
                   color: RGBColor = None, alignment=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def _add_header_row(table, texts: List[str]):
    """Format the first row of a table as a dark header."""
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        _set_cell_shading(cell, _BG_HEADER)
        _set_cell_text(cell, text, bold=True, size=9, color=_WHITE)


def _set_table_style(table):
    """Apply consistent table formatting."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set column widths to auto-fit
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.space_before = Pt(2)
                paragraph.space_after = Pt(2)


# =========================================================================
# Document section builders
# =========================================================================

def _generate_doc_id(report: ComparabilityReport) -> str:
    """Generate a document identifier from product name and timestamp."""
    ts = report.timestamp or datetime.datetime.now(datetime.timezone.utc).isoformat()
    date_part = ts[:10].replace("-", "")
    product_short = report.product_name.replace(" ", "-")[:20]
    return f"CMP-{product_short}-{date_part}"


def _add_cover_page(doc: Document, report: ComparabilityReport, doc_id: str):
    """Add cover page with title, product info, and date."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph("")

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Comparability Assessment Report")
    run.font.size = Pt(28)
    run.font.color.rgb = _DARK
    run.bold = True

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("ICH Q5E Compliance")
    run.font.size = Pt(16)
    run.font.color.rgb = _GRAY

    doc.add_paragraph("")

    # Product info block
    info_items = [
        ("Product", report.product_name),
        ("Manufacturing Change", report.change_description or "Not specified"),
        ("Document ID", doc_id),
        ("Date", (report.timestamp or "")[:10] or datetime.date.today().isoformat()),
        ("Assessment Verdict", report.overall_verdict),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f"{label}: ")
        run_label.font.size = Pt(12)
        run_label.font.color.rgb = _GRAY
        run_label.bold = True
        run_value = p.add_run(value)
        run_value.font.size = Pt(12)
        run_value.font.color.rgb = _DARK

    # Spacer and disclaimer
    doc.add_paragraph("")
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "This report was generated by the Biologics Decision Engine. "
        "All conclusions should be reviewed by qualified personnel."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = _GRAY
    run.italic = True

    doc.add_page_break()


def _add_executive_summary(doc: Document, report: ComparabilityReport):
    """Section 1: Executive Summary."""
    doc.add_heading("1. Executive Summary", level=1)

    # Verdict paragraph
    p = doc.add_paragraph()
    run = p.add_run("Overall Verdict: ")
    run.bold = True
    run.font.size = Pt(12)
    verdict_run = p.add_run(report.overall_verdict)
    verdict_run.font.size = Pt(14)
    verdict_run.bold = True
    verdict_run.font.color.rgb = _verdict_color(report.overall_verdict)

    # Evidence strength
    p = doc.add_paragraph()
    run = p.add_run("Evidence Strength Index: ")
    run.bold = True
    p.add_run(f"{_get_esi(report):.1%}")

    # Key statistics
    doc.add_paragraph(
        f"This assessment evaluated {report.n_attributes} quality attributes "
        f"across {_count_categories(report)} analytical categories. "
        f"Of these, {report.n_cqa} were classified as Critical Quality Attributes (CQAs)."
    )

    # Key findings bullets
    doc.add_paragraph("Key Findings:", style="List Bullet")
    doc.add_paragraph(
        f"{report.n_comparable} of {report.n_attributes} attributes "
        f"demonstrated comparability ({_pct(report.n_comparable, report.n_attributes)})",
        style="List Bullet 2",
    )
    if report.n_flagged > 0:
        doc.add_paragraph(
            f"{report.n_flagged} attribute(s) flagged with major or critical concerns",
            style="List Bullet 2",
        )
    gaps_count = len(report.evidence_gaps)
    if gaps_count > 0:
        doc.add_paragraph(
            f"{gaps_count} evidence gap(s) identified requiring follow-up",
            style="List Bullet 2",
        )

    high_unc = report.uncertainty_summary.get("n_high_uncertainty", 0)
    if high_unc > 0:
        doc.add_paragraph(
            f"{high_unc} attribute(s) with high residual uncertainty (>0.5)",
            style="List Bullet 2",
        )


def _add_scope_and_background(doc: Document, report: ComparabilityReport):
    """Section 2: Scope & Background."""
    doc.add_heading("2. Scope & Background", level=1)

    doc.add_heading("2.1 Manufacturing Change Description", level=2)
    doc.add_paragraph(
        report.change_description or "No change description provided."
    )

    doc.add_heading("2.2 Regulatory Basis", level=2)
    doc.add_paragraph(
        "This assessment follows ICH Q5E: Comparability of Biotechnological/"
        "Biological Products Subject to Changes in their Manufacturing Process. "
        "The guideline establishes principles for evaluating the impact of "
        "manufacturing process changes on the quality, safety, and efficacy "
        "of biotechnological/biological products."
    )
    doc.add_paragraph(
        "Key regulatory references: ICH Q5E (2004), ICH Q6B Specifications "
        "for Biotechnological/Biological Products, ICH Q5C Stability Testing, "
        "and relevant FDA/EMA guidance documents on comparability and "
        "biosimilarity assessments."
    )

    doc.add_heading("2.3 Assessment Methodology", level=2)
    doc.add_paragraph(
        "The comparability assessment follows a structured, attribute-by-attribute "
        "approach comprising six integrated analytical stages:"
    )
    steps = [
        "Data Harmonization: Unit normalization and field mapping across pre- and post-change datasets.",
        "CQA Classification: Risk-based designation of attributes as CQA, KQA, QA, or Monitor using RPN scoring.",
        "Comparability Scoring: Quantitative delta analysis with category-specific tolerances and weighted aggregation.",
        "Uncertainty Assessment: Five-dimensional residual uncertainty quantification per attribute.",
        "Evidence Closure: Gap identification and prioritization based on severity and regulatory impact.",
        "Verdict Generation: Confidence-weighted overall determination with per-attribute action recommendations.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")


def _add_cqa_summary(doc: Document, report: ComparabilityReport):
    """Section 3: CQA Summary Table."""
    doc.add_heading("3. Critical Quality Attribute Summary", level=1)

    if not report.cqa_summary:
        doc.add_paragraph("No CQA classification data available.")
        return

    doc.add_paragraph(
        "Quality attributes were classified using a risk-based approach "
        "incorporating impact, detectability, and controllability scores "
        "to compute a Risk Priority Number (RPN). Attributes are designated "
        "as CQA (Critical), KQA (Key), QA (Quality Attribute), or Monitor."
    )

    headers = ["Attribute", "Category", "RPN Score", "Designation", "Rationale"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _add_header_row(table, headers)

    for cqa in report.cqa_summary:
        row = table.add_row()
        _set_cell_text(row.cells[0], cqa.get("name", ""))
        _set_cell_text(row.cells[1], cqa.get("category", "").title())
        _set_cell_text(row.cells[2], str(cqa.get("rpn", "")),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        designation = cqa.get("designation", "")
        _set_cell_text(row.cells[3], designation, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if designation == "CQA":
            _set_cell_shading(row.cells[3], _BG_RED)
        elif designation == "KQA":
            _set_cell_shading(row.cells[3], _BG_YELLOW)
        _set_cell_text(row.cells[4], cqa.get("rationale", ""), size=8)

    _set_table_style(table)


def _add_attribute_comparison(doc: Document, report: ComparabilityReport):
    """Section 4: Attribute-by-Attribute Comparison."""
    doc.add_heading("4. Attribute-by-Attribute Comparison", level=1)

    if not report.attribute_results:
        doc.add_paragraph("No attribute data available.")
        return

    doc.add_paragraph(
        "The following table presents the head-to-head comparison of "
        "pre-change and post-change analytical results for each quality attribute. "
        "Cells are color-coded by concern level: green (none/minor), "
        "yellow (major), red (critical)."
    )

    headers = ["Attribute", "Category", "Pre-Change", "Post-Change",
               "Delta %", "Score", "Concern", "Action"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _add_header_row(table, headers)

    for ar in report.attribute_results:
        row = table.add_row()
        bg = _concern_color(ar.concern)

        # Attribute name
        _set_cell_text(row.cells[0], ar.name, bold=ar.is_cqa)
        if ar.is_cqa:
            # Mark CQAs with bold text
            pass  # already bold

        _set_cell_text(row.cells[1], ar.category.title())

        # Values with units
        pre_str = f"{ar.pre_value:.4g}" + (f" {ar.unit}" if ar.unit else "")
        post_str = f"{ar.post_value:.4g}" + (f" {ar.unit}" if ar.unit else "")
        _set_cell_text(row.cells[2], pre_str, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(row.cells[3], post_str, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

        # Delta
        _set_cell_text(row.cells[4], f"{ar.delta_pct:+.1f}%",
                       alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_shading(row.cells[4], bg)

        # Score
        _set_cell_text(row.cells[5], f"{ar.score:.3f}",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Concern
        _set_cell_text(row.cells[6], ar.concern.upper(), bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_shading(row.cells[6], bg)

        # Action
        action_level = ""
        if ar.action:
            action_level = ar.action.get("action_level", "")
        _set_cell_text(row.cells[7], action_level, bold=True,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
        if action_level:
            _set_cell_shading(row.cells[7], _concern_action_color(action_level))

    _set_table_style(table)


def _add_uncertainty_assessment(doc: Document, report: ComparabilityReport):
    """Section 5: Uncertainty Assessment."""
    doc.add_heading("5. Uncertainty Assessment", level=1)

    us = report.uncertainty_summary
    doc.add_paragraph(
        "Residual uncertainty is scored on a 0-1 scale across five dimensions: "
        "analytical method variability, lot-to-lot variability, sample size, "
        "functional correlation availability, and regulatory precedent. "
        "Higher values indicate greater uncertainty in the comparability conclusion."
    )

    # Summary stats
    p = doc.add_paragraph()
    run = p.add_run("Overall Uncertainty Summary")
    run.bold = True
    run.font.size = Pt(11)

    stats = [
        f"Mean uncertainty: {us.get('mean_uncertainty', 0):.3f}",
        f"Maximum uncertainty: {us.get('max_uncertainty', 0):.3f}",
        f"Attributes with high uncertainty (>0.5): {us.get('n_high_uncertainty', 0)}",
    ]
    for s in stats:
        doc.add_paragraph(s, style="List Bullet")

    # Per-attribute uncertainty table
    if report.attribute_results:
        doc.add_paragraph("")
        headers = ["Attribute", "Category", "CQA", "Uncertainty", "Level"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _add_header_row(table, headers)

        for ar in report.attribute_results:
            row = table.add_row()
            _set_cell_text(row.cells[0], ar.name)
            _set_cell_text(row.cells[1], ar.category.title())
            _set_cell_text(row.cells[2], ar.cqa_designation,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row.cells[3], f"{ar.uncertainty:.3f}",
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)

            # Level classification
            if ar.uncertainty > 0.7:
                level = "Very High"
                bg = _BG_RED
            elif ar.uncertainty > 0.5:
                level = "High"
                bg = _BG_YELLOW
            elif ar.uncertainty > 0.3:
                level = "Moderate"
                bg = _BG_LIGHT_GRAY
            else:
                level = "Low"
                bg = _BG_GREEN

            _set_cell_text(row.cells[4], level, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_shading(row.cells[4], bg)

        _set_table_style(table)

    # High uncertainty attributes detail
    high_attrs = us.get("high_uncertainty_attributes", [])
    if high_attrs:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Attributes Requiring Uncertainty Reduction")
        run.bold = True
        for name in high_attrs:
            doc.add_paragraph(
                f"{name}: Consider additional lots, orthogonal methods, "
                "or expanded replicates to reduce residual uncertainty.",
                style="List Bullet",
            )


def _add_evidence_gaps(doc: Document, report: ComparabilityReport):
    """Section 6: Evidence Gaps."""
    doc.add_heading("6. Evidence Gaps", level=1)

    if not report.evidence_gaps:
        doc.add_paragraph(
            "No significant evidence gaps were identified. The available data "
            "provides adequate coverage for the comparability conclusion."
        )
        return

    doc.add_paragraph(
        "The following evidence gaps were identified during the assessment. "
        "Each gap represents an area where additional data would strengthen "
        "the comparability conclusion or is required before proceeding."
    )

    for i, gap in enumerate(report.evidence_gaps, 1):
        doc.add_paragraph(f"{gap}", style="List Bullet")

    # Recommended additional studies
    doc.add_heading("6.1 Recommended Additional Studies", level=2)
    doc.add_paragraph(
        "To close the identified evidence gaps, the following studies "
        "are recommended, listed in order of priority:"
    )

    # Derive study recommendations from gaps and attribute actions
    study_recs = _derive_study_recommendations(report)
    if study_recs:
        for rec in study_recs:
            doc.add_paragraph(rec, style="List Bullet")
    else:
        doc.add_paragraph(
            "Specific study recommendations are provided in the "
            "Action Recommendations section below."
        )


def _add_action_recommendations(doc: Document, report: ComparabilityReport):
    """Section 7: Action Recommendations."""
    doc.add_heading("7. Action Recommendations", level=1)

    # Per-attribute actions
    doc.add_heading("7.1 Per-Attribute Actions", level=2)

    attrs_with_actions = [ar for ar in report.attribute_results if ar.action]
    if not attrs_with_actions:
        doc.add_paragraph("No per-attribute action data available.")
    else:
        headers = ["Attribute", "Action", "Rationale", "Next Evidence", "Effort"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _add_header_row(table, headers)

        for ar in attrs_with_actions:
            action = ar.action
            row = table.add_row()
            _set_cell_text(row.cells[0], ar.name, bold=ar.is_cqa)

            action_level = action.get("action_level", "")
            _set_cell_text(row.cells[1], action_level, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_shading(row.cells[1], _concern_action_color(action_level))

            _set_cell_text(row.cells[2], action.get("rationale", ""), size=8)
            _set_cell_text(row.cells[3], action.get("next_best_evidence", ""), size=8)
            _set_cell_text(row.cells[4], action.get("estimated_effort", ""), size=8)

        _set_table_style(table)

    # Overall action summary
    doc.add_heading("7.2 Overall Action Summary", level=2)

    if report.action_summary:
        summary = report.action_summary
        p = doc.add_paragraph()
        run = p.add_run("Overall Recommended Action: ")
        run.bold = True
        p.add_run(summary.get("overall_action", "N/A"))

        p = doc.add_paragraph()
        run = p.add_run("Regulatory Risk Level: ")
        run.bold = True
        p.add_run(summary.get("regulatory_risk", "N/A"))

        p = doc.add_paragraph()
        run = p.add_run("Estimated Timeline: ")
        run.bold = True
        p.add_run(summary.get("estimated_timeline", "N/A"))

        # Breakdown
        doc.add_paragraph(
            f"Action Breakdown: "
            f"{summary.get('n_proceed', 0)} PROCEED, "
            f"{summary.get('n_supplement', 0)} SUPPLEMENT, "
            f"{summary.get('n_monitor', 0)} MONITOR, "
            f"{summary.get('n_investigate', 0)} INVESTIGATE, "
            f"{summary.get('n_defer', 0)} DEFER"
        )

        # Critical attributes
        critical = summary.get("critical_attributes", [])
        if critical:
            p = doc.add_paragraph()
            run = p.add_run("Critical Attributes Requiring Immediate Attention: ")
            run.bold = True
            p.add_run(", ".join(critical))

        # Next steps
        next_steps = summary.get("next_steps", [])
        if next_steps:
            doc.add_paragraph("")
            p = doc.add_paragraph()
            run = p.add_run("Recommended Next Steps:")
            run.bold = True
            for step in next_steps:
                doc.add_paragraph(step, style="List Number")
    else:
        doc.add_paragraph("No overall action summary available.")

    # Regulatory references
    doc.add_heading("7.3 Regulatory References", level=2)
    _add_regulatory_references(doc, report)


def _add_judgment_core_section(doc: Document, report: ComparabilityReport):
    """Section 7.5: Judgment Core Analysis (Phase 4E, additive).

    Only rendered when judgment_core_verdict is populated (None-guarded).
    """
    if not getattr(report, 'judgment_core_verdict', None):
        return

    doc.add_heading("7.5 Judgment Core Analysis", level=2)

    # Verdict and confidence
    p = doc.add_paragraph()
    run = p.add_run("Judgment Core Verdict: ")
    run.bold = True
    p.add_run(report.judgment_core_verdict or "N/A")

    p = doc.add_paragraph()
    run = p.add_run("Confidence: ")
    run.bold = True
    conf = report.judgment_confidence
    band = report.judgment_confidence_band or ""
    p.add_run(f"{conf:.2f} ({band})" if conf is not None else "N/A")

    # Abstain flag
    if getattr(report, 'abstain_flag', None):
        p = doc.add_paragraph()
        run = p.add_run("ABSTAIN: ")
        run.bold = True
        run.font.color.rgb = _RED
        p.add_run("System abstained from judgment due to unresolvable conflicts.")

    # Blocking clusters
    blocking = getattr(report, 'blocking_clusters', None)
    if blocking:
        doc.add_paragraph(f"{len(blocking)} blocking cluster(s) identified:")
        for bc in blocking:
            doc.add_paragraph(
                f"{bc.get('category', 'Unknown')} -- {bc.get('risk_semantics', 'N/A')}: "
                f"{bc.get('reason', '')[:150]}",
                style="List Bullet",
            )

    # Decision rule IDs
    rule_ids = getattr(report, 'decision_rule_ids', None)
    if rule_ids:
        doc.add_paragraph(f"Decision rules applied: {', '.join(rule_ids)}")

    # What would change
    wwc = getattr(report, 'what_would_change_verdict', None)
    if wwc:
        doc.add_paragraph("Counterfactual analysis (what would change the verdict):")
        for cf in wwc:
            doc.add_paragraph(
                f"If {cf.get('current_gap', 'gap')} resolved: "
                f"verdict would become {cf.get('verdict_would_become', 'N/A')}",
                style="List Bullet",
            )


def _add_conclusion(doc: Document, report: ComparabilityReport):
    """Section 8: Conclusion."""
    doc.add_heading("8. Conclusion", level=1)

    # Verdict restatement
    p = doc.add_paragraph()
    run = p.add_run(f"Comparability Verdict: {report.overall_verdict}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = _verdict_color(report.overall_verdict)

    # Supporting rationale
    if report.overall_verdict == "Comparable":
        doc.add_paragraph(
            f"Based on the comprehensive attribute-by-attribute assessment, "
            f"{report.n_comparable} of {report.n_attributes} quality attributes "
            f"demonstrate comparability between pre-change and post-change material. "
            f"The evidence strength index of {_get_esi(report):.1%} supports "
            f"the conclusion that the manufacturing change does not adversely "
            f"impact the quality, safety, or efficacy of the product."
        )
    elif report.overall_verdict == "Not Comparable":
        flagged_names = [ar.name for ar in report.attribute_results
                         if ar.concern in ("major", "critical")]
        doc.add_paragraph(
            f"The assessment identified {report.n_flagged} attribute(s) with "
            f"significant concerns: {', '.join(flagged_names)}. "
            f"These findings indicate that comparability has not been "
            f"demonstrated for the manufacturing change as currently characterized. "
            f"Additional studies or process modifications are required before "
            f"a comparability conclusion can be supported."
        )
    else:
        doc.add_paragraph(
            f"The available evidence is insufficient to make a definitive "
            f"comparability determination. Key uncertainties remain in "
            f"{report.uncertainty_summary.get('n_high_uncertainty', 0)} "
            f"attribute(s). Additional batch data and/or analytical studies "
            f"are needed to support a conclusion."
        )

    # Conditions for proceeding
    doc.add_heading("8.1 Conditions for Proceeding", level=2)
    if report.overall_verdict == "Comparable":
        doc.add_paragraph(
            "Subject to standard quality oversight, the manufacturing change "
            "may proceed. The following conditions apply:"
        )
        doc.add_paragraph(
            "Maintain routine batch release testing and trending per existing protocols.",
            style="List Bullet",
        )
        monitor_attrs = [ar.name for ar in report.attribute_results
                         if ar.action and ar.action.get("action_level") == "MONITOR"]
        if monitor_attrs:
            doc.add_paragraph(
                f"Implement enhanced monitoring for: {', '.join(monitor_attrs)}.",
                style="List Bullet",
            )
        supplement_attrs = [ar.name for ar in report.attribute_results
                           if ar.action and ar.action.get("action_level") == "SUPPLEMENT"]
        if supplement_attrs:
            doc.add_paragraph(
                f"Collect supplementary data for: {', '.join(supplement_attrs)}.",
                style="List Bullet",
            )
    elif report.overall_verdict == "Not Comparable":
        doc.add_paragraph(
            "The manufacturing change should not proceed until the following "
            "conditions are met:"
        )
        for ar in report.attribute_results:
            if ar.concern in ("major", "critical") and ar.action:
                doc.add_paragraph(
                    f"{ar.name}: {ar.action.get('next_best_evidence', 'Investigate and resolve.')}",
                    style="List Bullet",
                )
    else:
        doc.add_paragraph(
            "Additional data collection is required before a determination "
            "can be made. See Section 6 (Evidence Gaps) for specific recommendations."
        )

    # Timeline
    doc.add_heading("8.2 Recommended Timeline", level=2)
    if report.action_summary and report.action_summary.get("estimated_timeline"):
        doc.add_paragraph(
            f"Estimated timeline to resolution: "
            f"{report.action_summary['estimated_timeline']}"
        )
    else:
        if report.overall_verdict == "Comparable":
            doc.add_paragraph("No additional timeline requirements. Standard batch release applies.")
        elif report.overall_verdict == "Not Comparable":
            doc.add_paragraph(
                "Estimated 2-4 months for root-cause investigation and supplementary studies."
            )
        else:
            doc.add_paragraph(
                "Estimated 1-3 months for additional data collection and re-assessment."
            )


# =========================================================================
# Helper functions
# =========================================================================

def _get_esi(report: ComparabilityReport) -> float:
    """Get evidence strength index, handling both field name variants."""
    return getattr(report, "evidence_strength_index",
                   getattr(report, "confidence", 0.0))


def _count_categories(report: ComparabilityReport) -> int:
    """Count unique attribute categories."""
    return len(set(ar.category for ar in report.attribute_results))


def _pct(numerator: int, denominator: int) -> str:
    """Format a percentage string."""
    if denominator == 0:
        return "N/A"
    return f"{100 * numerator / denominator:.0f}%"


def _derive_study_recommendations(report: ComparabilityReport) -> List[str]:
    """Extract concrete study recommendations from attribute actions."""
    recs = []
    seen = set()
    # Prioritize DEFER and INVESTIGATE actions
    for ar in report.attribute_results:
        if ar.action:
            level = ar.action.get("action_level", "")
            nbe = ar.action.get("next_best_evidence", "")
            if level in ("DEFER", "INVESTIGATE") and nbe and nbe not in seen:
                recs.append(f"{ar.name} ({level}): {nbe}")
                seen.add(nbe)
    for ar in report.attribute_results:
        if ar.action:
            level = ar.action.get("action_level", "")
            nbe = ar.action.get("next_best_evidence", "")
            if level in ("SUPPLEMENT", "MONITOR") and nbe and nbe not in seen:
                recs.append(f"{ar.name} ({level}): {nbe}")
                seen.add(nbe)
    return recs


def _add_regulatory_references(doc: Document, report: ComparabilityReport):
    """Add regulatory reference section based on categories assessed."""
    categories = set(ar.category for ar in report.attribute_results)

    _REFS = {
        "identity": "ICH Q5E Section 2.1 -- Identity characterization",
        "purity": "ICH Q5E Section 2.2 -- Purity, impurities, and contaminants; ICH Q6B",
        "potency": "ICH Q5E Section 2.3 -- Biological activity/potency; FDA Potency Guidance (2011)",
        "safety": "ICH Q5E Section 2.4 -- Immunochemical properties; ICH Q5A Viral Safety",
        "stability": "ICH Q5E Section 3 -- Stability; ICH Q1A/Q5C",
        "physicochemical": "ICH Q5E Section 2.2 -- Physicochemical properties; ICH Q6B",
    }

    for cat in sorted(categories):
        ref = _REFS.get(cat, f"ICH Q5E -- {cat.title()} attributes")
        doc.add_paragraph(ref, style="List Bullet")

    doc.add_paragraph(
        "ICH Q5E (2004): Comparability of Biotechnological/Biological Products "
        "Subject to Changes in Their Manufacturing Process",
        style="List Bullet",
    )


# =========================================================================
# Main generator
# =========================================================================

def generate_comparability_report(
    report: ComparabilityReport,
    output_path: str,
) -> str:
    """Generate an ICH Q5E-compliant comparability report as DOCX.

    Parameters
    ----------
    report : ComparabilityReport
        The structured comparability assessment output from the pipeline.
    output_path : str
        Path where the .docx file will be written.

    Returns
    -------
    str
        The absolute path to the generated report file.
    """
    doc = Document()

    # -- Document defaults --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = _DARK

    doc_id = _generate_doc_id(report)

    # -- Build sections --
    _add_cover_page(doc, report, doc_id)
    _add_executive_summary(doc, report)
    _add_scope_and_background(doc, report)
    _add_cqa_summary(doc, report)
    _add_attribute_comparison(doc, report)
    _add_uncertainty_assessment(doc, report)
    _add_evidence_gaps(doc, report)
    _add_action_recommendations(doc, report)
    _add_judgment_core_section(doc, report)
    _add_conclusion(doc, report)

    # -- Footer --
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"--- End of Report --- | {doc_id} | "
        f"Generated: {datetime.datetime.now(datetime.timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = _GRAY
    run.italic = True

    # -- Save --
    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)

    return output_path
