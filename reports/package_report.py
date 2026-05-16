"""
Package Assessment Report — DOCX export for multi-document CMC packages.

Generates a structured report with 7 sections:
1. Executive Summary
2. Document Inventory
3. Evidence Matrix
4. Gap Analysis
5. Cross-document Findings
6. Predicted Reviewer Questions
7. Appendix (extraction data)
"""

from __future__ import annotations

import os
from datetime import datetime
from typing import Any, Dict, List

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def generate_package_report(
    package_overview: Dict[str, Any],
    output_path: str,
) -> str:
    """Generate a DOCX package assessment report.

    Args:
        package_overview: Output from build_package_overview().
        output_path: Path to write the DOCX file.

    Returns:
        The output_path on success.
    """
    doc = Document()

    # Base styles
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    _add_title(doc, package_overview)
    _add_executive_summary(doc, package_overview)
    _add_document_inventory(doc, package_overview)
    _add_evidence_matrix(doc, package_overview)
    _add_gap_analysis(doc, package_overview)
    _add_cross_document_findings(doc, package_overview)
    _add_reviewer_questions(doc, package_overview)
    _add_appendix(doc, package_overview)
    _add_footer(doc)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _add_title(doc: Document, ov: Dict) -> None:
    title = doc.add_heading("CMC Package Assessment Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Package ID: {ov.get('package_id', 'N/A')}  |  "
        f"Date: {datetime.now().strftime('%Y-%m-%d')}  |  "
        f"Documents: {ov.get('n_documents', 0)}",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _add_executive_summary(doc: Document, ov: Dict) -> None:
    doc.add_heading("1. Executive Summary", level=1)

    verdict = ov.get("package_verdict_display", "Unknown")
    confidence = ov.get("package_confidence", 0)
    rationale = ov.get("package_rationale", "")

    p = doc.add_paragraph()
    run = p.add_run(f"Package Verdict: {verdict}")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Confidence: {confidence:.0%}")
    doc.add_paragraph(rationale)

    # Coverage summary
    coverage = ov.get("document_coverage", {})
    missing = ov.get("missing_types", [])
    if missing:
        doc.add_paragraph(
            f"Missing document types: {', '.join(t.replace('_', ' ').title() for t in missing)}",
        )
    else:
        doc.add_paragraph("All required document types are present.")

    doc.add_paragraph()


def _add_document_inventory(doc: Document, ov: Dict) -> None:
    doc.add_heading("2. Document Inventory", level=1)

    summaries = ov.get("document_summaries", [])
    if not summaries:
        doc.add_paragraph("No documents assessed.")
        return

    table = doc.add_table(rows=1 + len(summaries), cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Filename", "Document Type", "Verdict", "Posture", "Confidence"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for idx, ds in enumerate(summaries):
        row = table.rows[idx + 1]
        row.cells[0].text = ds.get("filename", "")
        row.cells[1].text = ds.get("doc_type", "")
        row.cells[2].text = ds.get("analytical_conclusion", "N/A")
        row.cells[3].text = ds.get("package_posture", "N/A")
        row.cells[4].text = f"{ds.get('confidence', 0):.0%}"
        if ds.get("error"):
            row.cells[2].text = "ERROR"
            row.cells[3].text = ds["error"][:50]

    doc.add_paragraph()


def _add_evidence_matrix(doc: Document, ov: Dict) -> None:
    doc.add_heading("3. Evidence Matrix", level=1)
    doc.add_paragraph(
        "Coverage of Critical Quality Attributes across submitted documents."
    )

    coverage = ov.get("document_coverage", {})
    summaries = ov.get("document_summaries", [])

    # CQA categories
    cqa_categories = [
        "Aggregation (HMW%)", "Charge Variants", "Glycosylation",
        "Potency", "Purity", "Identity/Structure",
    ]
    doc_types = [ds.get("doc_type", "") for ds in summaries]

    table = doc.add_table(rows=1 + len(cqa_categories), cols=1 + len(doc_types))
    table.style = "Table Grid"

    # Header row
    table.rows[0].cells[0].text = "CQA Category"
    for i, dt in enumerate(doc_types):
        table.rows[0].cells[i + 1].text = dt[:15]
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    # Matrix body — which CQAs each doc type typically covers
    _type_coverage = {
        "CHARACTERIZATION": {0, 1, 2, 3, 4, 5},
        "STABILITY": {0, 1, 4},
        "ANALYTICAL_METHOD": {0, 1, 4},
        "COMPARABILITY": {0, 1, 2, 3, 4},
    }
    for row_idx, cqa in enumerate(cqa_categories):
        table.rows[row_idx + 1].cells[0].text = cqa
        for col_idx, dt in enumerate(doc_types):
            covered = row_idx in _type_coverage.get(dt, set())
            table.rows[row_idx + 1].cells[col_idx + 1].text = "Yes" if covered else "--"

    doc.add_paragraph()


def _add_gap_analysis(doc: Document, ov: Dict) -> None:
    doc.add_heading("4. Gap Analysis", level=1)

    summaries = ov.get("document_summaries", [])
    has_gaps = False

    for ds in summaries:
        if ds.get("error"):
            doc.add_paragraph(
                f"{ds['filename']}: Processing error — {ds['error'][:100]}",
                style="List Bullet",
            )
            has_gaps = True
            continue

        posture = ds.get("package_posture", "")
        if posture not in ("Sufficient", "Complete", "Ready"):
            doc.add_paragraph(
                f"{ds['filename']} ({ds['doc_type']}): {ds['analytical_conclusion']} / {posture}",
                style="List Bullet",
            )
            has_gaps = True

    missing = ov.get("missing_types", [])
    for mt in missing:
        doc.add_paragraph(
            f"Missing: {mt.replace('_', ' ').title()} — required for complete CMC package",
            style="List Bullet",
        )
        has_gaps = True

    if not has_gaps:
        doc.add_paragraph("No significant gaps identified.")

    doc.add_paragraph()


def _add_cross_document_findings(doc: Document, ov: Dict) -> None:
    doc.add_heading("5. Cross-Document Findings", level=1)

    flags = ov.get("cross_document_flags", [])
    if not flags:
        doc.add_paragraph("No cross-document inconsistencies detected.")
    else:
        for f in flags:
            severity = f.get("severity", "info").upper()
            desc = f.get("description", "")
            doc.add_paragraph(f"[{severity}] {desc}", style="List Bullet")

    doc.add_paragraph()


def _add_reviewer_questions(doc: Document, ov: Dict) -> None:
    doc.add_heading("6. Predicted Reviewer Questions", level=1)

    questions = ov.get("reviewer_questions", [])
    if not questions:
        doc.add_paragraph("No reviewer questions predicted.")
        return

    for i, q in enumerate(questions, 1):
        source = q.get("source_doc_type", "")
        question = q.get("question", "")
        confidence = q.get("confidence", "moderate")
        doc.add_paragraph(
            f"{i}. [{source}] ({confidence.upper()}) {question}",
            style="List Number",
        )

    doc.add_paragraph()


def _add_appendix(doc: Document, ov: Dict) -> None:
    doc.add_heading("7. Appendix: Document Details", level=1)

    summaries = ov.get("document_summaries", [])
    for ds in summaries:
        doc.add_heading(f"{ds['filename']} — {ds['doc_type']}", level=2)
        doc.add_paragraph(f"Classification confidence: {ds.get('classification_confidence', 0):.0%}")
        doc.add_paragraph(f"Verdict: {ds.get('analytical_conclusion', 'N/A')} / {ds.get('package_posture', 'N/A')}")
        doc.add_paragraph(f"Assessment confidence: {ds.get('confidence', 0):.0%}")
        if ds.get("error"):
            doc.add_paragraph(f"Error: {ds['error']}")

    doc.add_paragraph()


def _add_footer(doc: Document) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph(
        "For decision support only. Not regulatory advice. "
        "Verify all findings with source documents and a qualified regulatory professional."
    )
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(100, 116, 139)
