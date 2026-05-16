"""
B5: Generate a synthetic characterization report DOCX for testing.

Run once to create tests/fixtures/test_characterization_report.docx.
This script is also imported by test_characterization_extractor.py
to create the fixture on-the-fly if needed.
"""

import os
import sys

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, PROJECT_ROOT)


def create_characterization_report_docx(path: str) -> None:
    """Create a synthetic ICH Q6B characterization report DOCX.

    Includes sections and tables covering all 8 required Q6B areas.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Characterization Report -- mAb-Test IgG1", 0)
    doc.add_paragraph(
        "This document presents the physicochemical and biological characterization "
        "of mAb-Test IgG1 (anti-IL6R monoclonal antibody) per ICH Q6B guidelines. "
        "The reference standard lot RS-2026-001 was used throughout."
    )
    doc.add_paragraph(
        "Reference Standard Lot: RS-2026-001 (USP reference)"
    )

    # ---- 1. Primary Structure ----
    doc.add_heading("1. Primary Structure", level=1)
    doc.add_paragraph(
        "The primary structure of mAb-Test IgG1 was confirmed by peptide mapping "
        "using LC-MS/MS. The amino acid sequence showed 100% coverage. "
        "N-terminal sequencing was performed by Edman degradation."
    )
    table = doc.add_table(rows=4, cols=4, style="Table Grid")
    headers = ["Parameter", "Method", "Result", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Molecular Weight (intact)", "LC-MS", "148235", "Da"],
        ["Sequence Coverage", "Peptide Mapping LC-MS/MS", "100", "%"],
        ["N-terminal Sequence", "Edman Degradation", "Confirmed", ""],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 2. Higher-Order Structure ----
    doc.add_heading("2. Higher-Order Structure", level=1)
    doc.add_paragraph(
        "Higher-order structure (HOS) analysis was performed using circular "
        "dichroism (CD) spectroscopy (far-UV and near-UV), differential scanning "
        "calorimetry (DSC), and FTIR spectroscopy. Secondary structure content "
        "was consistent with an IgG1 antibody."
    )
    table = doc.add_table(rows=4, cols=4, style="Table Grid")
    headers = ["Assay", "Method", "Result", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Secondary Structure (beta-sheet)", "Far-UV CD", "68.5", "%"],
        ["Thermal Stability (Tm1)", "DSC", "71.2", "degC"],
        ["Structural Similarity", "FTIR", "0.98", "correlation"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 3. Aggregation / Size Variants ----
    doc.add_heading("3. Aggregation and Size Variants", level=1)
    doc.add_paragraph(
        "Size variant analysis was performed by SEC-HPLC, AUC (analytical "
        "ultracentrifugation), and DLS (dynamic light scattering). "
        "HMW species (aggregates) were quantified."
    )
    table = doc.add_table(rows=5, cols=4, style="Table Grid")
    headers = ["Attribute", "Method", "Value", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Monomer", "SEC-HPLC", "98.2", "%"],
        ["HMW (Aggregates)", "SEC-HPLC", "1.2", "%"],
        ["LMW (Fragments)", "SEC-HPLC", "0.6", "%"],
        ["Hydrodynamic Radius", "DLS", "5.4", "nm"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 4. Charge Heterogeneity ----
    doc.add_heading("4. Charge Heterogeneity", level=1)
    doc.add_paragraph(
        "Charge variant analysis was performed by cation exchange "
        "chromatography (CEX) and icIEF (imaged capillary isoelectric focusing)."
    )
    table = doc.add_table(rows=4, cols=4, style="Table Grid")
    headers = ["Attribute", "Method", "Value", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Acidic Variants", "CEX", "22.1", "%"],
        ["Main Charge Peak", "CEX", "58.3", "%"],
        ["Basic Variants", "CEX", "19.6", "%"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 5. Glycosylation / PTMs ----
    doc.add_heading("5. Glycosylation and Post-Translational Modifications", level=1)
    doc.add_paragraph(
        "N-glycan profiling was performed by HILIC-MS following PNGase F "
        "enzymatic release. Post-translational modifications including "
        "deamidation, oxidation, and C-terminal lysine clipping were assessed."
    )
    table = doc.add_table(rows=6, cols=4, style="Table Grid")
    headers = ["Parameter", "Method", "Value", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["G0F", "N-glycan HILIC-MS", "42.3", "%"],
        ["G1F", "N-glycan HILIC-MS", "28.7", "%"],
        ["G2F", "N-glycan HILIC-MS", "8.1", "%"],
        ["Afucosylation", "N-glycan HILIC-MS", "4.2", "%"],
        ["High Mannose (Man5)", "N-glycan HILIC-MS", "2.8", "%"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 6. Biological Activity / Potency ----
    doc.add_heading("6. Biological Activity and Potency", level=1)
    doc.add_paragraph(
        "Biological activity was assessed using a cell-based reporter gene "
        "assay measuring IL-6 receptor inhibition. ADCC activity was also "
        "evaluated. Relative potency was determined against the reference standard."
    )
    table = doc.add_table(rows=4, cols=4, style="Table Grid")
    headers = ["Assay", "Method", "Value", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Relative Potency", "Cell-based Reporter Gene", "102.5", "%"],
        ["ADCC Activity", "ADCC Reporter Assay", "98.7", "%"],
        ["Binding EC50", "Cell-based Assay", "0.45", "nM"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 7. Immunochemical Properties ----
    doc.add_heading("7. Immunochemical Properties", level=1)
    doc.add_paragraph(
        "Target binding was assessed by ELISA and surface plasmon resonance "
        "(SPR) using a Biacore instrument. FcRn binding was confirmed."
    )
    table = doc.add_table(rows=4, cols=4, style="Table Grid")
    headers = ["Parameter", "Method", "Value", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Target Binding (Kd)", "SPR / Biacore", "0.12", "nM"],
        ["Target Binding", "ELISA", "Positive", ""],
        ["FcRn Binding (Kd)", "SPR", "1.8", "uM"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 8. Purity / Impurities ----
    doc.add_heading("8. Purity and Impurities", level=1)
    doc.add_paragraph(
        "Purity was assessed by rCE-SDS (reduced and non-reduced) and "
        "RP-HPLC. Process-related impurities including host cell protein (HCP) "
        "and residual DNA were quantified."
    )
    table = doc.add_table(rows=6, cols=4, style="Table Grid")
    headers = ["Attribute", "Method", "Value", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Purity (reduced)", "rCE-SDS", "98.5", "%"],
        ["Purity (non-reduced)", "CE-SDS", "97.8", "%"],
        ["Purity", "RP-HPLC", "99.1", "%"],
        ["HCP", "ELISA", "12.5", "ppm"],
        ["Residual DNA", "qPCR", "0.8", "pg/mg"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- Summary ----
    doc.add_heading("9. Summary", level=1)
    doc.add_paragraph(
        "The characterization study demonstrates that mAb-Test IgG1 has the "
        "expected physicochemical and biological properties consistent with an "
        "IgG1 monoclonal antibody. All critical quality attributes meet "
        "acceptance criteria. Potency: 102.5% relative to reference standard "
        "lot RS-2026-001."
    )

    doc.save(path)


if __name__ == "__main__":
    fixture_path = os.path.join(
        PROJECT_ROOT, "tests", "fixtures", "test_characterization_report.docx"
    )
    os.makedirs(os.path.dirname(fixture_path), exist_ok=True)
    create_characterization_report_docx(fixture_path)
    print(f"Created: {fixture_path}")
