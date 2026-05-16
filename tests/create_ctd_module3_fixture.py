"""
Phase 4F: Create synthetic CTD Module 3 test fixture.

Generates a DOCX file with sections 3.2.S.1 through 3.2.S.7
and 3.2.P.1 through 3.2.P.8 containing representative content
and tables for each section.
"""

import os
import sys

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)


def create_ctd_module3_fixture():
    """Create a synthetic CTD Module 3 DOCX document."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title
    title = doc.add_heading("CTD Module 3 - Quality", level=0)

    # ---------------------------------------------------------------
    # 3.2.S Drug Substance sections
    # ---------------------------------------------------------------

    # 3.2.S.1 General Information
    doc.add_heading("3.2.S.1 General Information (Drug Substance)", level=1)
    doc.add_paragraph(
        "The drug substance is a monoclonal antibody (mAb) produced by recombinant "
        "DNA technology in Chinese Hamster Ovary (CHO) cells. The INN is Examplemab."
    )
    doc.add_paragraph(
        "Molecular formula: Approximately 145 kDa IgG1 monoclonal antibody. "
        "The molecule targets the CD20 antigen on B lymphocytes."
    )

    # 3.2.S.2 Manufacture
    doc.add_heading("3.2.S.2 Manufacture (Drug Substance)", level=1)
    doc.add_paragraph(
        "The drug substance is manufactured at BioFacility Corp., Site A. "
        "The manufacturing process consists of upstream cell culture, harvest, "
        "and downstream purification including Protein A chromatography, "
        "ion exchange, and viral inactivation steps."
    )
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ["Parameter", "Target", "Actual"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Cell Culture Duration", "14 days", "14 days"],
        ["Harvest Titer", ">2.0 g/L", "2.8 g/L"],
        ["Protein A Yield", ">80%", "92%"],
    ]
    for r, row_data in enumerate(data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    # 3.2.S.3 Characterisation
    doc.add_heading("3.2.S.3 Characterisation (Drug Substance)", level=1)
    doc.add_paragraph(
        "Comprehensive characterization was performed per ICH Q6B including "
        "primary structure confirmation by peptide mapping (LC-MS/MS), "
        "higher-order structure by CD and DSC, glycan profiling by HILIC-MS, "
        "charge variants by icIEF, and size variants by SEC-HPLC."
    )
    doc.add_paragraph(
        "Reference standard lot: RS-2024-001. "
        "HMW: 1.2%. Main charge peak: 65.3%. Afucosylation: 8.5%. "
        "Relative potency: 98%."
    )

    # Characterization table
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    headers = ["Attribute", "Method", "Result", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    char_data = [
        ["HMW (Aggregates)", "SEC-HPLC", "1.2", "%"],
        ["Monomer Purity", "SEC-HPLC", "97.5", "%"],
        ["Main Charge Peak", "icIEF", "65.3", "%"],
        ["Afucosylation", "HILIC-MS", "8.5", "%"],
        ["Relative Potency", "Cell-based assay", "98", "%"],
        ["Binding Affinity (Kd)", "SPR/Biacore", "0.15", "nM"],
    ]
    for r, row_data in enumerate(char_data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    # 3.2.S.4 Control of Drug Substance
    doc.add_heading("3.2.S.4 Control of Drug Substance", level=1)
    doc.add_paragraph(
        "Method validation was performed per ICH Q2(R2) for all release testing methods. "
        "Specificity, linearity, accuracy, precision (repeatability and intermediate), "
        "range, LOD, LOQ, and robustness were assessed."
    )
    doc.add_paragraph(
        "SE-HPLC Method Validation: Recovery: 99.5%. RSD: 1.2%. R2: 0.9998. "
        "LOD: 0.05%. LOQ: 0.1%."
    )

    # Method validation table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ["Validation Parameter", "Result", "Acceptance Criteria"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    method_data = [
        ["Accuracy (Recovery)", "99.5%", "98.0-102.0%"],
        ["Precision (%RSD)", "1.2%", "<=2.0%"],
        ["Linearity (R2)", "0.9998", ">=0.999"],
        ["LOQ", "0.1%", "Report value"],
    ]
    for r, row_data in enumerate(method_data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    # 3.2.S.5 Reference Standards
    doc.add_heading("3.2.S.5 Reference Standards (Drug Substance)", level=1)
    doc.add_paragraph(
        "The primary reference standard (lot RS-2024-001) was fully characterized "
        "and qualified. Traceability to WHO International Standard established."
    )

    # 3.2.S.6 Container Closure System
    doc.add_heading("3.2.S.6 Container Closure System (Drug Substance)", level=1)
    doc.add_paragraph(
        "Drug substance is stored in 2L polycarbonate bottles with HDPE screw caps. "
        "Container closure integrity testing per USP <1207> confirmed adequate seal."
    )

    # 3.2.S.7 Stability (Drug Substance)
    doc.add_heading("3.2.S.7 Stability (Drug Substance)", level=1)
    doc.add_paragraph(
        "Stability studies were conducted per ICH Q1A/Q5C at long-term (5C), "
        "accelerated (25C/60%RH), and stress (40C/75%RH) conditions. "
        "Proposed shelf-life of 24 months at 2-8 degrees C."
    )

    # Stability table
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Table Grid'
    headers = ["Attribute", "Condition", "T=0", "3M", "6M", "12M"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    stab_data = [
        ["Purity (SEC)", "5C", "97.5", "97.3", "97.1", "96.8"],
        ["HMW", "5C", "1.2", "1.3", "1.4", "1.6"],
        ["Potency", "5C", "98", "97", "96", "95"],
        ["Purity (SEC)", "40C/75RH", "97.5", "95.0", "92.1", "88.5"],
    ]
    for r, row_data in enumerate(stab_data, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    # ---------------------------------------------------------------
    # 3.2.P Drug Product sections
    # ---------------------------------------------------------------

    # 3.2.P.1 Description and Composition
    doc.add_heading("3.2.P.1 Description and Composition (Drug Product)", level=1)
    doc.add_paragraph(
        "The drug product is a sterile, preservative-free solution for IV infusion. "
        "Each vial contains 100 mg of Examplemab in 10 mL (10 mg/mL) formulated "
        "in 20 mM histidine, 240 mM trehalose, 0.02% polysorbate 80, pH 6.0."
    )

    # 3.2.P.2 Pharmaceutical Development
    doc.add_heading("3.2.P.2 Pharmaceutical Development", level=1)
    doc.add_paragraph(
        "Formulation development studies optimized buffer, pH, surfactant, and "
        "tonicity agent to maximize stability. Forced degradation studies confirmed "
        "the stability-indicating nature of all analytical methods."
    )

    # 3.2.P.3 Manufacture
    doc.add_heading("3.2.P.3 Manufacture (Drug Product)", level=1)
    doc.add_paragraph(
        "Drug product manufacture includes thawing of drug substance, dilution, "
        "sterile filtration (0.22 um), aseptic fill into glass vials, "
        "stoppering, and capping. Process validation completed for 3 consecutive batches."
    )

    # 3.2.P.4 Control of Excipients
    doc.add_heading("3.2.P.4 Control of Excipients", level=1)
    doc.add_paragraph(
        "All excipients comply with USP/NF or Ph. Eur. monographs. "
        "Polysorbate 80 is compendial grade. Histidine and trehalose meet USP specifications."
    )

    # 3.2.P.5 Control of Drug Product
    doc.add_heading("3.2.P.5 Control of Drug Product", level=1)
    doc.add_paragraph(
        "Release testing specifications and validated methods for drug product. "
        "Method validation per ICH Q2(R2). Specificity, linearity, accuracy, "
        "precision, and robustness demonstrated for all methods."
    )

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ["Test", "Method", "Acceptance Criteria"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    dp_spec = [
        ["Appearance", "Visual", "Clear, colorless to slightly yellow"],
        ["pH", "Potentiometry", "5.7 - 6.3"],
        ["Protein Concentration", "UV A280", "9.0 - 11.0 mg/mL"],
        ["Purity (SEC)", "SE-HPLC", ">= 95.0%"],
        ["Potency", "Cell-based assay", "80 - 120%"],
    ]
    for r, row_data in enumerate(dp_spec, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    # 3.2.P.6 Reference Standards
    doc.add_heading("3.2.P.6 Reference Standards (Drug Product)", level=1)
    doc.add_paragraph(
        "The same primary reference standard (lot RS-2024-001) is used for "
        "drug product release testing."
    )

    # 3.2.P.7 Container Closure System
    doc.add_heading("3.2.P.7 Container Closure System (Drug Product)", level=1)
    doc.add_paragraph(
        "20 mL Type I borosilicate glass vials with 20 mm fluoropolymer-coated "
        "butyl rubber stoppers and aluminum flip-off seals. "
        "Container closure integrity verified per USP <1207>."
    )

    # 3.2.P.8 Stability (Drug Product)
    doc.add_heading("3.2.P.8 Stability (Drug Product)", level=1)
    doc.add_paragraph(
        "Drug product stability studies per ICH Q1A/Q5C. "
        "Long-term (5C), accelerated (25C/60%RH), and stress (40C/75%RH). "
        "Proposed shelf-life of 24 months at 2-8 degrees C."
    )

    # DP Stability table
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    headers = ["Attribute", "Condition", "T=0", "3M", "6M", "12M"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    dp_stab = [
        ["Purity (SEC)", "5C", "97.2", "97.0", "96.8", "96.5"],
        ["Potency", "5C", "100", "99", "98", "97"],
        ["Subvisible Particles", "5C", "12", "15", "18", "22"],
    ]
    for r, row_data in enumerate(dp_stab, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    # Save
    output_path = os.path.join(
        os.path.dirname(__file__), "fixtures", "test_ctd_module3.docx"
    )
    doc.save(output_path)
    print(f"Created CTD Module 3 fixture: {output_path}")
    return output_path


if __name__ == "__main__":
    create_ctd_module3_fixture()
