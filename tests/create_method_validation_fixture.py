"""
D3: Generate a synthetic method validation report DOCX for testing.

Creates tests/fixtures/test_method_validation_report.docx with ICH Q2(R2)
validation tables covering accuracy, precision, linearity, etc.
"""

import os
import sys

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, PROJECT_ROOT)


def create_method_validation_report_docx(path: str) -> None:
    """Create a synthetic ICH Q2(R2) method validation report DOCX.

    Includes validation summary tables for all major ICH Q2 parameters.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Analytical Method Validation Report", 0)
    doc.add_paragraph(
        "Method Validation Report for SEC-HPLC Purity Method. "
        "This report describes the validation of the SEC-HPLC method "
        "for determination of monomer purity and size variants (HMW, LMW) "
        "of mAb-Val-01 per ICH Q2(R2) guidelines."
    )
    doc.add_paragraph(
        "Analytical Method: SEC-HPLC for Size Variant Analysis"
    )
    doc.add_paragraph(
        "Method validation of SEC-HPLC method. "
        "Method name: SEC-HPLC Purity"
    )

    # ---- 1. Specificity ----
    doc.add_heading("1. Specificity", level=1)
    doc.add_paragraph(
        "Specificity was demonstrated by analyzing placebo, reference standard, "
        "and stressed samples. No interference from the placebo matrix was observed."
    )
    table = doc.add_table(rows=4, cols=3, style="Table Grid")
    headers = ["Parameter", "Result", "Criteria"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Placebo interference", "None detected", "No interference"],
        ["Peak resolution (monomer/HMW)", "2.8", ">= 1.5"],
        ["Selectivity", "Demonstrated", "Demonstrated"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 2. Linearity ----
    doc.add_heading("2. Linearity", level=1)
    doc.add_paragraph(
        "Linearity was assessed over the range of 0.1 to 2.0 mg/mL "
        "(50-200% of nominal concentration). "
        "R2 = 0.9998. Correlation coefficient: 0.9999."
    )
    table = doc.add_table(rows=4, cols=3, style="Table Grid")
    headers = ["Parameter", "Result", "Criteria"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["R2 (correlation coefficient)", "0.9998", ">= 0.999"],
        ["Slope", "1.002", "0.95-1.05"],
        ["Y-intercept", "0.15", "< 2.0%"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 3. Range ----
    doc.add_heading("3. Range", level=1)
    doc.add_paragraph(
        "The analytical range was confirmed as 0.1 to 2.0 mg/mL, "
        "corresponding to 50-200% of the nominal working concentration."
    )

    # ---- 4. Accuracy ----
    doc.add_heading("4. Accuracy", level=1)
    doc.add_paragraph(
        "Accuracy was determined by spiked recovery experiments at three "
        "concentration levels (80%, 100%, 120%). "
        "Mean recovery: 100.2%. Recovery: 100.2%."
    )
    table = doc.add_table(rows=4, cols=3, style="Table Grid")
    headers = ["Parameter", "Result", "Criteria"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Recovery at 80%", "99.5%", "98-102%"],
        ["Recovery at 100%", "100.2%", "98-102%"],
        ["Recovery at 120%", "100.8%", "98-102%"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 5. Precision (Repeatability) ----
    doc.add_heading("5. Precision -- Repeatability", level=1)
    doc.add_paragraph(
        "Repeatability (intra-day precision) was assessed by 6 replicate "
        "injections of the reference standard. %RSD = 0.8%."
    )
    table = doc.add_table(rows=3, cols=3, style="Table Grid")
    headers = ["Parameter", "Result", "Criteria"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["RSD (n=6)", "0.8%", "<= 2.0%"],
        ["Precision (intra-assay)", "0.8%", "<= 2.0%"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 6. Precision (Intermediate) ----
    doc.add_heading("6. Precision -- Intermediate Precision", level=1)
    doc.add_paragraph(
        "Intermediate precision was evaluated across two analysts and "
        "three days. Inter-day %RSD = 1.2%."
    )
    table = doc.add_table(rows=3, cols=3, style="Table Grid")
    headers = ["Parameter", "Result", "Criteria"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Inter-day RSD", "1.2%", "<= 3.0%"],
        ["Inter-analyst RSD", "1.5%", "<= 3.0%"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 7. LOD / LOQ ----
    doc.add_heading("7. LOD and LOQ", level=1)
    doc.add_paragraph(
        "The LOD (limit of detection) was determined to be 0.02 mg/mL. "
        "The LOQ (limit of quantitation) was determined to be 0.05 mg/mL. "
        "These are adequate for detection and quantitation of impurity species."
    )
    table = doc.add_table(rows=3, cols=3, style="Table Grid")
    headers = ["Parameter", "Result", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["LOD", "0.02", "mg/mL"],
        ["LOQ", "0.05", "mg/mL"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- 8. Robustness ----
    doc.add_heading("8. Robustness", level=1)
    doc.add_paragraph(
        "Robustness was evaluated by deliberate variation of column temperature "
        "(+/- 2 degC), flow rate (+/- 10%), and mobile phase composition (+/- 2%). "
        "The method was found to be robust under all tested conditions."
    )

    # ---- 9. System Suitability ----
    doc.add_heading("9. System Suitability", level=1)
    doc.add_paragraph(
        "System suitability criteria were met in all validation runs. "
        "Theoretical plates > 5000, tailing factor < 1.5."
    )

    # ---- Summary ----
    doc.add_heading("10. Summary", level=1)
    doc.add_paragraph(
        "The SEC-HPLC purity method has been validated per ICH Q2(R2). "
        "All validation parameters meet acceptance criteria. "
        "The method is suitable for its intended use in release and stability testing."
    )

    doc.save(path)


if __name__ == "__main__":
    fixture_path = os.path.join(
        PROJECT_ROOT, "tests", "fixtures", "test_method_validation_report.docx"
    )
    os.makedirs(os.path.dirname(fixture_path), exist_ok=True)
    create_method_validation_report_docx(fixture_path)
    print(f"Created: {fixture_path}")
