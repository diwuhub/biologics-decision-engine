"""
C3: Generate a synthetic stability report DOCX for testing.

Creates tests/fixtures/test_stability_report.docx with timepoint tables,
storage conditions, and OOS flag scenarios.
"""

import os
import sys

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, PROJECT_ROOT)


def create_stability_report_docx(path: str) -> None:
    """Create a synthetic ICH Q1A/Q5C stability report DOCX.

    Includes timepoint tables with multiple conditions and an OOS flag.
    """
    from docx import Document

    doc = Document()
    doc.add_heading("Stability Study Report -- mAb-Stab-01", 0)
    doc.add_paragraph(
        "This stability study was conducted per ICH Q1A(R2) and ICH Q5C "
        "guidelines to support the proposed shelf-life of 24 months at "
        "2-8 degrees C for mAb-Stab-01 drug product."
    )
    doc.add_paragraph(
        "Shelf-life of 24 months is proposed based on long-term, accelerated, "
        "and stress stability data."
    )

    # ---- Long-term stability (5C) ----
    doc.add_heading("1. Long-Term Stability (5 degC)", level=1)
    doc.add_paragraph(
        "Long-term stability study at 5 degC (refrigerated storage). "
        "Samples were tested at T=0, 3M, 6M, 9M, 12M, 18M, and 24M."
    )
    table = doc.add_table(rows=6, cols=8, style="Table Grid")
    headers = ["Attribute", "T=0", "3M", "6M", "9M", "12M", "18M", "24M"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Purity (SEC)", "99.2", "99.1", "99.0", "98.8", "98.6", "98.3", "98.0"],
        ["HMW (%)", "0.5", "0.6", "0.7", "0.8", "0.9", "1.1", "1.3"],
        ["Potency (%)", "102", "101", "101", "100", "100", "99", "98"],
        ["pH", "6.2", "6.2", "6.2", "6.1", "6.1", "6.1", "6.1"],
        ["Subvisible Particles", "12", "15", "18", "20", "22", "25", "28"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- Accelerated stability (25C/60%RH) ----
    doc.add_heading("2. Accelerated Stability (25 degC / 60% RH)", level=1)
    doc.add_paragraph(
        "Accelerated stability study at 25 degC / 60% RH per ICH Q1A."
    )
    table = doc.add_table(rows=5, cols=5, style="Table Grid")
    headers = ["Attribute", "T=0", "3M", "6M", "12M"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Purity (SEC)", "99.2", "98.5", "97.8", "96.5"],
        ["HMW (%)", "0.5", "1.0", "1.8", "2.8"],
        ["Potency (%)", "102", "99", "96", "92"],
        ["Appearance", "Clear", "Clear", "Clear", "Slightly opalescent"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- Stress stability (40C/75%RH) ----
    doc.add_heading("3. Stress Stability (40 degC / 75% RH)", level=1)
    doc.add_paragraph(
        "Stress stability study at 40 degC / 75% RH. "
        "OOS observed for potency at 3M timepoint."
    )
    table = doc.add_table(rows=4, cols=4, style="Table Grid")
    headers = ["Attribute", "T=0", "1M", "3M"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["Purity (SEC)", "99.2", "96.1", "91.5"],
        ["HMW (%)", "0.5", "2.5", "5.8"],
        ["Potency (%)", "102", "88", "FAIL - OOS"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    # ---- Summary ----
    doc.add_heading("4. Summary and Conclusions", level=1)
    doc.add_paragraph(
        "The long-term stability data at 5 degC supports a shelf-life of 24 months. "
        "Accelerated data at 25 degC / 60% RH shows acceptable degradation kinetics. "
        "Stress data at 40 degC / 75% RH demonstrates expected rapid degradation "
        "with an OOS event for potency at 3 months. "
        "All stability-indicating attributes monitored include purity, aggregation, "
        "potency, pH, and subvisible particles."
    )

    doc.save(path)


if __name__ == "__main__":
    fixture_path = os.path.join(
        PROJECT_ROOT, "tests", "fixtures", "test_stability_report.docx"
    )
    os.makedirs(os.path.dirname(fixture_path), exist_ok=True)
    create_stability_report_docx(fixture_path)
    print(f"Created: {fixture_path}")
