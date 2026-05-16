"""
Tests for the ICH Q5E Comparability Report Generator.

Verifies that:
  1. A DOCX report can be generated from the demo case
  2. The output file exists and is a valid DOCX
  3. The pipeline integration (generate_report flag) works
  4. Edge cases (empty report) are handled gracefully
"""

import json
import os
import sys
import tempfile

import pytest

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, PROJECT_ROOT)

from pipelines.comparability import run_comparability_assessment
from pipelines.schemas import ComparabilityReport
from reports.comparability_report import generate_comparability_report


DEMO_PATH = os.path.join(PROJECT_ROOT, "benchmarks", "cases", "COMP-001.json")
FALLBACK_DEMO_PATH = os.path.join(PROJECT_ROOT, "benchmarks", "mab_process_change_case.json")


def _get_demo_path():
    if os.path.exists(DEMO_PATH):
        return DEMO_PATH
    return FALLBACK_DEMO_PATH


def _load_demo_case():
    with open(_get_demo_path()) as f:
        return json.load(f)


# =========================================================================
# Test 1: Generate report from demo case -- file exists and is valid DOCX
# =========================================================================

def test_generate_report_from_demo_case():
    case = _load_demo_case()
    report = run_comparability_assessment(
        pre_change_data=case,
        product_name=case.get("product_name", "TestProduct"),
        change_description=case.get("change_description", "Test change"),
    )

    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = os.path.join(tmpdir, "test_report.docx")
        result_path = generate_comparability_report(report, output_path)

        # File exists
        assert os.path.exists(result_path), f"Report file not created at {result_path}"

        # File is non-empty
        file_size = os.path.getsize(result_path)
        assert file_size > 1000, f"Report file too small ({file_size} bytes)"

        # File is a valid DOCX (ZIP with correct magic bytes)
        with open(result_path, "rb") as f:
            magic = f.read(4)
        assert magic == b"PK\x03\x04", "File is not a valid DOCX/ZIP archive"

        # File can be opened by python-docx
        from docx import Document
        doc = Document(result_path)
        assert len(doc.paragraphs) > 10, "Report has too few paragraphs"
        assert len(doc.tables) > 0, "Report has no tables"


# =========================================================================
# Test 2: Report contains expected content
# =========================================================================

def test_report_contains_expected_content():
    case = _load_demo_case()
    report = run_comparability_assessment(
        pre_change_data=case,
        product_name=case.get("product_name", "TestProduct"),
        change_description=case.get("change_description", "Test change"),
    )

    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = os.path.join(tmpdir, "test_report.docx")
        generate_comparability_report(report, output_path)

        from docx import Document
        doc = Document(output_path)

        # Collect all text
        all_text = "\n".join(p.text for p in doc.paragraphs)

        # Check for key sections
        assert "Comparability Assessment Report" in all_text
        assert "Executive Summary" in all_text
        assert "Scope" in all_text
        assert "Critical Quality Attribute" in all_text
        assert "Attribute-by-Attribute" in all_text
        assert "Uncertainty" in all_text
        assert "Evidence Gaps" in all_text
        assert "Action Recommendations" in all_text
        assert "Conclusion" in all_text

        # Check for product name
        product_name = case.get("product_name", "TestProduct")
        assert product_name in all_text

        # Check for verdict
        assert report.overall_verdict in all_text


# =========================================================================
# Test 3: Pipeline integration -- generate_report flag
# =========================================================================

def test_pipeline_integration_generate_report():
    case = _load_demo_case()

    with tempfile.TemporaryDirectory() as tmpdir:
        report_path = os.path.join(tmpdir, "pipeline_report.docx")
        report = run_comparability_assessment(
            pre_change_data=case,
            product_name=case.get("product_name", "TestProduct"),
            change_description=case.get("change_description", "Test change"),
            generate_report=True,
            report_path=report_path,
        )

        # Pipeline still returns the report object
        assert isinstance(report, ComparabilityReport)
        assert report.n_attributes > 0

        # DOCX was generated
        assert os.path.exists(report_path), "Pipeline did not generate report file"
        assert os.path.getsize(report_path) > 1000


# =========================================================================
# Test 4: Empty report handles gracefully
# =========================================================================

def test_empty_report_generates_valid_docx():
    report = run_comparability_assessment(
        pre_change_data={"attributes": []},
        product_name="EmptyTest",
        change_description="No data",
    )

    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = os.path.join(tmpdir, "empty_report.docx")
        result_path = generate_comparability_report(report, output_path)

        assert os.path.exists(result_path)
        assert os.path.getsize(result_path) > 500

        from docx import Document
        doc = Document(result_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Insufficient Evidence" in all_text


# =========================================================================
# Test 5: All COMP benchmark cases generate valid reports
# =========================================================================

def test_all_comp_cases_generate_reports():
    cases_dir = os.path.join(PROJECT_ROOT, "benchmarks", "cases")
    if not os.path.isdir(cases_dir):
        pytest.skip("benchmarks/cases directory not found")

    comp_files = [f for f in os.listdir(cases_dir)
                  if f.startswith("COMP-") and f.endswith(".json")]
    assert len(comp_files) > 0, "No COMP case files found"

    with tempfile.TemporaryDirectory() as tmpdir:
        for case_file in comp_files:
            with open(os.path.join(cases_dir, case_file)) as f:
                case = json.load(f)

            report = run_comparability_assessment(
                pre_change_data=case,
                product_name=case.get("product_name", "Product"),
                change_description=case.get("change_description", ""),
            )

            output_path = os.path.join(tmpdir, f"{case_file.replace('.json', '')}_report.docx")
            result_path = generate_comparability_report(report, output_path)
            assert os.path.exists(result_path), f"Failed to generate report for {case_file}"
            assert os.path.getsize(result_path) > 500, f"Report too small for {case_file}"
