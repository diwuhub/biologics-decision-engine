"""
v5.2 Release Gate: Document-Driven End-to-End Tests.

Two DOCX files are created programmatically, ingested through the full
ingest_docx() -> run_comparability_assessment pipeline, and validated for
enriched output behavior.

Case 1: Well-structured mAb comparability report (happy path, scale-up)
Case 2: Spec-rich ADC site transfer with method context

Run: python3 -m pytest tests/test_docx_e2e.py -v --tb=short
"""

from __future__ import annotations

import os
import sys
import tempfile

import pytest

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, PROJECT_ROOT)

FIXTURES_DIR = os.path.join(PROJECT_ROOT, "tests", "fixtures")


def _create_happy_path_docx(path: str) -> None:
    """Create a well-structured comparability report DOCX."""
    from docx import Document

    doc = Document()
    doc.add_heading("Comparability Assessment Report", 0)
    doc.add_heading("Product: mAb-DocTest-01", level=1)
    doc.add_paragraph(
        "This comparability study evaluates the impact of a manufacturing process change "
        "(scale-up from 2000L to 5000L bioreactor) on the monoclonal antibody mAb-DocTest-01. "
        "The post-approval commercial product was assessed across physicochemical, purity, "
        "potency, and stability attributes."
    )
    doc.add_heading("Comparability Data", level=1)

    table = doc.add_table(rows=7, cols=6, style="Table Grid")
    headers = ["Attribute", "Category", "Pre-Change", "Post-Change", "Specification", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["SEC Monomer", "Purity", "98.5", "98.2", ">=95.0", "%"],
        ["SEC Aggregate", "Purity", "0.8", "1.0", "<=2.0", "%"],
        ["CEX Main Peak", "Physicochemical", "65.2", "64.8", "55.0-75.0", "%"],
        ["Potency (Bioassay)", "Potency", "102.0", "99.5", "80.0-120.0", "%RP"],
        ["Osmolality", "Physicochemical", "290.0", "292.0", "270.0-330.0", "mOsm/kg"],
        ["Endotoxin", "Purity", "0.05", "0.04", "<=0.5", "EU/mL"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "All tested attributes met acceptance criteria. The manufacturing scale-up "
        "from 2000L to 5000L did not adversely impact product quality."
    )
    doc.save(path)


def _create_specrich_docx(path: str) -> None:
    """Create a spec-rich ADC site transfer report DOCX."""
    from docx import Document

    doc = Document()
    doc.add_heading("Comparability Protocol CP-2024-042", 0)
    doc.add_heading("Product: BioFusion-ADC-07", level=1)
    doc.add_paragraph(
        "Antibody-drug conjugate BioFusion-ADC-07 comparability study following "
        "a site transfer from Facility A (Dublin) to Facility B (Singapore). "
        "This post-approval change requires demonstration of comparability per "
        "ICH Q5E. The analytical method panel includes validated HPLC, CE-SDS, "
        "and cell-based potency assays."
    )

    doc.add_heading("Change Description", level=1)
    doc.add_paragraph(
        "The manufacturing site transfer involves identical equipment specifications "
        "but different water systems and facility utilities. Process parameters were "
        "locked; only site-specific factors differ."
    )

    doc.add_heading("Analytical Results", level=1)

    table = doc.add_table(rows=8, cols=6, style="Table Grid")
    headers = ["Test", "Category", "Pre-Change", "Post-Change", "Acceptance Criteria", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["DAR (Drug-Antibody Ratio)", "Physicochemical", "3.8", "3.9", "3.5-4.5", "ratio"],
        ["Free Drug", "Purity", "1.2", "1.5", "<=3.0", "%"],
        ["SEC Monomer", "Purity", "97.8", "97.1", ">=95.0", "%"],
        ["CE-SDS (NR) Main Band", "Purity", "92.5", "91.8", ">=85.0", "%"],
        ["Potency (Cell-based)", "Potency", "105.0", "98.0", "80.0-120.0", "%RP"],
        ["pH", "Physicochemical", "6.2", "6.1", "5.8-6.5", ""],
        ["Particulate (>=10um)", "Physicochemical", "150.0", "180.0", "<=600", "per container"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "BioFusion-ADC-07 manufactured at Facility B demonstrates comparability to "
        "Facility A material. All critical quality attributes meet acceptance criteria."
    )
    doc.save(path)


# =========================================================================
# Test class
# =========================================================================


class TestDocxEndToEnd:
    """Document-driven end-to-end tests: DOCX upload -> full pipeline -> enriched output."""

    @pytest.fixture(autouse=True)
    def setup_fixtures(self, tmp_path):
        """Create DOCX files in a temp directory."""
        self.happy_path = str(tmp_path / "test_comparability_report.docx")
        self.specrich_path = str(tmp_path / "test_specrich_method_report.docx")
        _create_happy_path_docx(self.happy_path)
        _create_specrich_docx(self.specrich_path)

    # ----- Case 1: Happy path -----

    def test_case1_ingestion_extracts_attributes(self):
        """DOCX ingestion extracts all 6 attributes from the comparability table."""
        from ingestion import ingest_docx

        result = ingest_docx(self.happy_path)
        assert result.n_attributes_extracted == 6
        assert result.n_tables_found == 1
        assert len(result.issues) == 0

    def test_case1_context_detection(self):
        """Context extractor detects mAb molecule class and process change type."""
        from ingestion import ingest_docx

        result = ingest_docx(self.happy_path)
        assert result.case_context.molecule_class == "mAb"
        assert "process" in result.case_context.change_type.lower() or \
               "scale" in result.case_context.change_type.lower()

    def test_case1_spec_limits_pass_through(self):
        """Spec limits from DOCX are passed through to pipeline input."""
        from ingestion import ingest_docx

        result = ingest_docx(self.happy_path)
        attrs = result.pipeline_input["attributes"]
        # SEC Monomer should have spec_lower=95.0
        sec_monomer = [a for a in attrs if "Monomer" in a["name"]]
        assert len(sec_monomer) == 1
        assert sec_monomer[0].get("spec_lower") == 95.0

    def test_case1_pipeline_verdict_comparable(self):
        """Happy path with all-within-spec attributes produces 'Comparable' verdict."""
        from ingestion import ingest_docx
        from pipelines.comparability import run_comparability_assessment

        result = ingest_docx(self.happy_path)
        report = run_comparability_assessment(
            pre_change_data=result.pipeline_input,
            product_name="mAb-DocTest-01",
            change_description=result.case_context.change_description,
        )
        assert report.overall_verdict == "Comparable"
        assert report.n_comparable == report.n_attributes

    def test_case1_enriched_output_has_actions(self):
        """Pipeline output includes action recommendations for each attribute."""
        from ingestion import ingest_docx
        from pipelines.comparability import run_comparability_assessment

        result = ingest_docx(self.happy_path)
        report = run_comparability_assessment(
            pre_change_data=result.pipeline_input,
            product_name="mAb-DocTest-01",
            change_description=result.case_context.change_description,
        )
        for ar in report.attribute_results:
            assert ar.action is not None, f"No action for {ar.name}"
            assert "action_level" in ar.action

    # ----- Case 2: Spec-rich ADC -----

    def test_case2_ingestion_extracts_attributes(self):
        """DOCX ingestion extracts all 7 attributes from the ADC table."""
        from ingestion import ingest_docx

        result = ingest_docx(self.specrich_path)
        assert result.n_attributes_extracted == 7
        assert result.n_tables_found == 1
        assert len(result.issues) == 0

    def test_case2_context_detects_adc(self):
        """Context extractor detects ADC molecule class and site transfer."""
        from ingestion import ingest_docx

        result = ingest_docx(self.specrich_path)
        assert result.case_context.molecule_class == "ADC"
        assert "site" in result.case_context.change_type.lower()

    def test_case2_spec_limits_range(self):
        """Range-style spec limits (e.g. '3.5-4.5') are parsed correctly."""
        from ingestion import ingest_docx

        result = ingest_docx(self.specrich_path)
        attrs = result.pipeline_input["attributes"]
        dar = [a for a in attrs if "DAR" in a["name"]]
        assert len(dar) == 1
        assert dar[0].get("spec_lower") == 3.5
        assert dar[0].get("spec_upper") == 4.5

    def test_case2_pipeline_produces_verdict(self):
        """ADC case runs through pipeline and produces a valid verdict."""
        from ingestion import ingest_docx
        from pipelines.comparability import run_comparability_assessment

        result = ingest_docx(self.specrich_path)
        report = run_comparability_assessment(
            pre_change_data=result.pipeline_input,
            product_name="BioFusion-ADC-07",
            change_description=result.case_context.change_description,
        )
        assert report.overall_verdict in (
            "Comparable", "Not Comparable", "Comparable with Conditions"
        )
        assert report.n_attributes == 7

    def test_case2_enriched_output_has_evidence(self):
        """Pipeline output includes evidence strength index."""
        from ingestion import ingest_docx
        from pipelines.comparability import run_comparability_assessment

        result = ingest_docx(self.specrich_path)
        report = run_comparability_assessment(
            pre_change_data=result.pipeline_input,
            product_name="BioFusion-ADC-07",
            change_description=result.case_context.change_description,
        )
        assert report.evidence_strength_index > 0.0
        # Each attribute should have action recommendations
        for ar in report.attribute_results:
            assert ar.action is not None, f"No action for {ar.name}"

    def test_case2_enriched_output_differentiated(self):
        """Enriched output shows richer reasoning than just pass/fail."""
        from ingestion import ingest_docx
        from pipelines.comparability import run_comparability_assessment

        result = ingest_docx(self.specrich_path)
        report = run_comparability_assessment(
            pre_change_data=result.pipeline_input,
            product_name="BioFusion-ADC-07",
            change_description=result.case_context.change_description,
        )
        # Check for narrative reasoning in attribute results
        reasoning_found = False
        for ar in report.attribute_results:
            if hasattr(ar, "posture_rationale") and ar.posture_rationale:
                reasoning_found = True
                break
            if ar.action and ar.action.get("rationale"):
                reasoning_found = True
                break
        assert reasoning_found, "No enriched narrative reasoning found in output"
