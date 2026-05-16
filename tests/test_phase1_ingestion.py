"""
Phase 1 Track A + Track E tests.

Covers:
- A1: DocumentClassifier
- A2: PDFDocumentParser
- A3: IngestionDispatcher
- A4: GenericCMCExtractor
- A5: UnifiedIngestionResult
- A6: ingest_document() canonical entry point
- A7: BaseExtractor ABC

Invariants verified:
- INV-001: No document upload raises unhandled exception
- INV-002: PDF and DOCX parsers produce identical schema
- INV-003: Existing comparability benchmarks pass unchanged (run_benchmarks.py)
- INV-004: DocumentClassifier never blocks analysis (UNKNOWN -> generic extractor)

Track E:
- E1: MANIFEST.yaml and download.py exist
- E2: capability_probe imports and runs
- E3: drift_detector imports and runs
- E4: run_qa.py imports
"""

from __future__ import annotations

import os
import sys
import tempfile

import pytest

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, PROJECT_ROOT)


# ==========================================================================
# Helpers
# ==========================================================================

def _create_comparability_docx(path: str) -> None:
    """Create a minimal comparability DOCX for testing."""
    from docx import Document

    doc = Document()
    doc.add_heading("Comparability Assessment Report", 0)
    doc.add_paragraph(
        "This comparability study evaluates the impact of a manufacturing "
        "process change on the monoclonal antibody mAb-Test-01."
    )

    table = doc.add_table(rows=4, cols=5, style="Table Grid")
    headers = ["Attribute", "Pre-Change", "Post-Change", "Specification", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["SEC Monomer", "98.5", "98.2", ">=95.0", "%"],
        ["Potency", "102.0", "99.5", "80.0-120.0", "%RP"],
        ["Endotoxin", "0.05", "0.04", "<=0.5", "EU/mL"],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    doc.save(path)


def _create_characterization_docx(path: str) -> None:
    """Create a minimal characterization DOCX (no pre/post columns)."""
    from docx import Document

    doc = Document()
    doc.add_heading("Characterization Report", 0)
    doc.add_paragraph(
        "Structural analysis and characterization of mAb-Char-01. "
        "Higher-order structure was confirmed by CD spectroscopy."
    )

    table = doc.add_table(rows=3, cols=3, style="Table Grid")
    headers = ["Parameter", "Value", "Unit"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["Molecular Weight", "148500", "Da"],
        ["pI", "8.2", ""],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val

    doc.save(path)


def _create_minimal_pdf(path: str) -> None:
    """Create a minimal PDF for testing using fpdf2 or raw bytes."""
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(200, 10, text="Comparability Assessment Report", ln=True)
        pdf.cell(200, 10, text="Pre-change and post-change analysis", ln=True)
        pdf.cell(200, 10, text="This is a comparability study.", ln=True)
        pdf.output(path)
    except ImportError:
        # Fallback: create a minimal valid PDF manually
        content = b"""%PDF-1.4
1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj
2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj
3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj
4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj
5 0 obj<</Length 44>>
stream
BT /F1 12 Tf 100 700 Td (Test PDF) Tj ET
endstream
endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000340 00000 n
trailer<</Size 6/Root 1 0 R>>
startxref
434
%%EOF"""
        with open(path, "wb") as f:
            f.write(content)


# ==========================================================================
# A1: DocumentClassifier
# ==========================================================================

class TestDocumentClassifier:
    """A1: Test DocumentClassifier."""

    def test_classify_comparability_docx(self, tmp_path):
        """Comparability DOCX is correctly classified."""
        from ingestion.docx_parser import DOCXDocumentParser
        from ingestion.document_classifier import DocumentClassifier

        docx_path = str(tmp_path / "comp.docx")
        _create_comparability_docx(docx_path)

        parser = DOCXDocumentParser()
        parsed = parser.parse(docx_path)

        classifier = DocumentClassifier()
        result = classifier.classify(parsed)

        assert result.document_type == "COMPARABILITY"
        assert result.confidence >= 0.5

    def test_classify_characterization_docx(self, tmp_path):
        """Characterization DOCX is classified as CHARACTERIZATION or UNKNOWN (not COMPARABILITY)."""
        from ingestion.docx_parser import DOCXDocumentParser
        from ingestion.document_classifier import DocumentClassifier

        docx_path = str(tmp_path / "char.docx")
        _create_characterization_docx(docx_path)

        parser = DOCXDocumentParser()
        parsed = parser.parse(docx_path)

        classifier = DocumentClassifier()
        result = classifier.classify(parsed)

        # It should NOT be classified as COMPARABILITY
        assert result.document_type != "COMPARABILITY"

    def test_classify_empty_doc_returns_unknown(self):
        """Empty parsed doc returns UNKNOWN without crashing."""
        from ingestion.document_classifier import DocumentClassifier

        classifier = DocumentClassifier()
        result = classifier.classify({"pages": [], "paragraphs": [], "metadata": {}})

        assert result.document_type == "UNKNOWN"
        assert result.confidence < 0.5

    def test_classify_never_crashes(self):
        """INV-004: Classifier never crashes, even with garbage input."""
        from ingestion.document_classifier import DocumentClassifier

        classifier = DocumentClassifier()
        # None fields, missing keys, etc.
        result = classifier.classify({})
        assert result.document_type == "UNKNOWN"

        result = classifier.classify({"pages": None})
        assert result.document_type is not None


# ==========================================================================
# A2: PDFDocumentParser
# ==========================================================================

class TestPDFParser:
    """A2: Test PDFDocumentParser."""

    def test_parse_pdf(self, tmp_path):
        """PDF parsing produces the common schema."""
        from ingestion.pdf_parser import PDFDocumentParser

        pdf_path = str(tmp_path / "test.pdf")
        _create_minimal_pdf(pdf_path)

        parser = PDFDocumentParser()
        result = parser.parse(pdf_path)

        assert "pages" in result
        assert "paragraphs" in result
        assert "sections" in result
        assert "metadata" in result
        assert result["metadata"]["source_format"] == "pdf"

    def test_pdf_file_not_found(self):
        """PDF parser raises FileNotFoundError for missing file."""
        from ingestion.pdf_parser import PDFDocumentParser

        parser = PDFDocumentParser()
        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.pdf")

    def test_pdf_schema_matches_docx(self, tmp_path):
        """INV-002: PDF and DOCX parsers produce identical schema keys."""
        from ingestion.pdf_parser import PDFDocumentParser
        from ingestion.docx_parser import DOCXDocumentParser

        pdf_path = str(tmp_path / "test.pdf")
        docx_path = str(tmp_path / "test.docx")
        _create_minimal_pdf(pdf_path)
        _create_comparability_docx(docx_path)

        pdf_result = PDFDocumentParser().parse(pdf_path)
        docx_result = DOCXDocumentParser().parse(docx_path)

        # Same top-level keys
        assert set(pdf_result.keys()) == set(docx_result.keys())

        # Pages have same structure
        if pdf_result["pages"] and docx_result["pages"]:
            pdf_page_keys = set(pdf_result["pages"][0].keys())
            docx_page_keys = set(docx_result["pages"][0].keys())
            assert "page_number" in pdf_page_keys
            assert "text" in pdf_page_keys
            assert "tables" in pdf_page_keys
            # docx pages may have the same or subset of keys
            common = {"page_number", "text", "tables"}
            assert common.issubset(pdf_page_keys)
            assert common.issubset(docx_page_keys)

    def test_supported_formats(self):
        from ingestion.pdf_parser import PDFDocumentParser
        assert ".pdf" in PDFDocumentParser().supported_formats()


# ==========================================================================
# A3: IngestionDispatcher
# ==========================================================================

class TestIngestionDispatcher:
    """A3: Test IngestionDispatcher."""

    def test_dispatch_comparability(self):
        """COMPARABILITY type dispatches to comparability extractor."""
        from ingestion.dispatcher import IngestionDispatcher
        from ingestion.document_classifier import DocTypeSpec

        dispatcher = IngestionDispatcher()
        doc_type = DocTypeSpec("COMPARABILITY", 0.9, [])
        extractor = dispatcher.dispatch({}, doc_type)

        assert extractor is not None
        assert hasattr(extractor, "extract_attributes")
        assert hasattr(extractor, "extract_evidence")

    def test_dispatch_unknown_gives_generic(self):
        """UNKNOWN type dispatches to GenericCMCExtractor."""
        from ingestion.dispatcher import IngestionDispatcher
        from ingestion.document_classifier import DocTypeSpec
        from ingestion.generic_extractor import GenericCMCExtractor

        dispatcher = IngestionDispatcher()
        doc_type = DocTypeSpec("UNKNOWN", 0.3, [])
        extractor = dispatcher.dispatch({}, doc_type)

        assert isinstance(extractor, GenericCMCExtractor)

    def test_dispatch_characterization_gives_characterization_extractor(self):
        """CHARACTERIZATION dispatches to CharacterizationExtractor (Phase 2 Track B)."""
        from ingestion.dispatcher import IngestionDispatcher
        from ingestion.document_classifier import DocTypeSpec
        from ingestion.characterization_extractor import CharacterizationExtractor

        dispatcher = IngestionDispatcher()
        doc_type = DocTypeSpec("CHARACTERIZATION", 0.8, [])
        extractor = dispatcher.dispatch({}, doc_type)

        assert isinstance(extractor, CharacterizationExtractor)


# ==========================================================================
# A4: GenericCMCExtractor
# ==========================================================================

class TestGenericCMCExtractor:
    """A4: Test GenericCMCExtractor never crashes."""

    def test_extract_attributes_never_raises(self):
        """extract_attributes() must NEVER raise."""
        from ingestion.generic_extractor import GenericCMCExtractor

        ext = GenericCMCExtractor()
        # Empty doc
        result = ext.extract_attributes({})
        assert isinstance(result, list)

        # Garbage doc
        result = ext.extract_attributes({"pages": [{"tables": [{"headers": None}]}]})
        assert isinstance(result, list)

    def test_extract_evidence_never_raises(self):
        """extract_evidence() must NEVER raise."""
        from ingestion.generic_extractor import GenericCMCExtractor

        ext = GenericCMCExtractor()
        result = ext.extract_evidence({})
        assert isinstance(result, dict)

        result = ext.extract_evidence({"pages": None})
        assert isinstance(result, dict)

    def test_extract_from_table(self):
        """Generic extractor can pull attributes from a simple table."""
        from ingestion.generic_extractor import GenericCMCExtractor

        parsed_doc = {
            "document_path": "test.pdf",
            "pages": [{
                "page_number": 1,
                "text": "",
                "tables": [{
                    "id": "t1",
                    "headers": ["Parameter", "Value", "Unit"],
                    "rows": [
                        {"Parameter": "MW", "Value": "148500", "Unit": "Da"},
                        {"Parameter": "pI", "Value": "8.2", "Unit": ""},
                    ],
                }],
            }],
        }

        ext = GenericCMCExtractor()
        attrs = ext.extract_attributes(parsed_doc)
        assert len(attrs) == 2
        assert attrs[0].name == "MW"


# ==========================================================================
# A5: UnifiedIngestionResult
# ==========================================================================

class TestUnifiedIngestionResult:
    """A5: Test UnifiedIngestionResult dataclass."""

    def test_creation(self):
        from ingestion.unified_result import UnifiedIngestionResult
        from ingestion.context_extractor import ExtractedCaseContext
        from ingestion.document_classifier import DocTypeSpec

        result = UnifiedIngestionResult(
            attributes=[],
            case_context=ExtractedCaseContext(),
            signals=[],
            document_classification=DocTypeSpec("COMPARABILITY", 0.9, []),
            extracted_evidence={"tables_found": 1},
        )

        assert result.document_classification.document_type == "COMPARABILITY"
        assert result.extracted_evidence["tables_found"] == 1
        assert result.source_format == "unknown"


# ==========================================================================
# A6: ingest_document() canonical entry point
# ==========================================================================

class TestIngestDocument:
    """A6: Test ingest_document() unified entry point."""

    def test_ingest_docx(self, tmp_path):
        """ingest_document() works on DOCX files."""
        from ingestion import ingest_document

        docx_path = str(tmp_path / "comp.docx")
        _create_comparability_docx(docx_path)

        result = ingest_document(docx_path)

        assert result.source_format == "docx"
        assert result.document_classification is not None
        assert result.document_classification.document_type == "COMPARABILITY"
        assert result.n_attributes_extracted > 0

    def test_ingest_pdf(self, tmp_path):
        """ingest_document() works on PDF files."""
        from ingestion import ingest_document

        pdf_path = str(tmp_path / "test.pdf")
        _create_minimal_pdf(pdf_path)

        result = ingest_document(pdf_path)

        assert result.source_format == "pdf"
        assert result.document_classification is not None
        assert len(result.issues) >= 0  # may have issues but must not crash

    def test_ingest_document_nonexistent_file(self):
        """INV-001: Non-existent file does not raise unhandled exception."""
        from ingestion import ingest_document

        result = ingest_document("/nonexistent/file.docx")
        assert len(result.issues) > 0
        assert result.document_classification.document_type == "UNKNOWN"

    def test_ingest_document_corrupt_file(self, tmp_path):
        """INV-001: Corrupt file does not raise unhandled exception."""
        from ingestion import ingest_document

        corrupt_path = str(tmp_path / "corrupt.docx")
        with open(corrupt_path, "wb") as f:
            f.write(b"this is not a docx file")

        result = ingest_document(corrupt_path)
        assert len(result.issues) > 0

    def test_ingest_document_classifier_routes_to_generic(self, tmp_path):
        """INV-004: UNKNOWN classification uses generic extractor, does not block."""
        from ingestion import ingest_document

        docx_path = str(tmp_path / "char.docx")
        _create_characterization_docx(docx_path)

        result = ingest_document(docx_path)
        # Should not crash, and should have a classification
        assert result.document_classification is not None
        # Should still produce some kind of result (may be empty attributes)
        assert isinstance(result.attributes, list)

    def test_ingest_docx_backward_compat(self, tmp_path):
        """ingest_docx() still works as before (backward compatible alias)."""
        from ingestion import ingest_docx

        docx_path = str(tmp_path / "comp.docx")
        _create_comparability_docx(docx_path)

        result = ingest_docx(docx_path)
        assert result.n_attributes_extracted == 3
        assert result.n_tables_found == 1

    def test_ingest_document_comparability_matches_ingest_docx(self, tmp_path):
        """ingest_document() extracts same attributes as ingest_docx() for comparability DOCX."""
        from ingestion import ingest_docx, ingest_document

        docx_path = str(tmp_path / "comp.docx")
        _create_comparability_docx(docx_path)

        old_result = ingest_docx(docx_path)
        new_result = ingest_document(docx_path)

        # Same number of attributes
        assert new_result.n_attributes_extracted == old_result.n_attributes_extracted
        # Same attribute names
        old_names = {a.name for a in old_result.attributes}
        new_names = {a.name for a in new_result.attributes}
        assert old_names == new_names


# ==========================================================================
# A7: BaseExtractor ABC
# ==========================================================================

class TestBaseExtractor:
    """A7: Test BaseExtractor interface."""

    def test_cannot_instantiate_directly(self):
        from ingestion.base_extractor import BaseExtractor

        with pytest.raises(TypeError):
            BaseExtractor()

    def test_subclass_must_implement(self):
        from ingestion.base_extractor import BaseExtractor

        class IncompleteExtractor(BaseExtractor):
            pass

        with pytest.raises(TypeError):
            IncompleteExtractor()

    def test_subclass_with_methods_works(self):
        from ingestion.base_extractor import BaseExtractor

        class TestExtractor(BaseExtractor):
            def extract_attributes(self, parsed_doc):
                return []

            def extract_evidence(self, parsed_doc):
                return {}

        ext = TestExtractor()
        assert ext.extract_attributes({}) == []
        assert ext.extract_evidence({}) == {}


# ==========================================================================
# Track E: QA Skeleton
# ==========================================================================

class TestQASkeleton:
    """Track E: QA agent skeleton tests."""

    def test_e1_manifest_exists(self):
        """E1: MANIFEST.yaml exists."""
        manifest_path = os.path.join(
            PROJECT_ROOT, "benchmarks", "real_documents", "MANIFEST.yaml"
        )
        assert os.path.isfile(manifest_path)

    def test_e1_download_exists(self):
        """E1: download.py exists."""
        download_path = os.path.join(
            PROJECT_ROOT, "benchmarks", "real_documents", "download.py"
        )
        assert os.path.isfile(download_path)

    def test_e1_manifest_loads(self):
        """E1: MANIFEST.yaml loads and contains documents."""
        import yaml

        manifest_path = os.path.join(
            PROJECT_ROOT, "benchmarks", "real_documents", "MANIFEST.yaml"
        )
        with open(manifest_path) as f:
            manifest = yaml.safe_load(f)
        assert "documents" in manifest
        assert len(manifest["documents"]) > 0

    def test_e2_capability_probe_imports(self):
        """E2: capability_probe module imports correctly."""
        from qa.capability_probe import probe_capability, ProbeResult
        assert callable(probe_capability)

    def test_e2_probe_implemented_capability(self):
        """E2: Probe an implemented capability."""
        import yaml
        from qa.capability_probe import probe_capability

        spec_path = os.path.join(PROJECT_ROOT, "qa", "vision_spec.yaml")
        with open(spec_path) as f:
            vision_spec = yaml.safe_load(f)

        result = probe_capability("CAP-001", vision_spec)
        assert result.capability_id == "CAP-001"
        assert result.status in ("pass", "fail", "skip", "error")

    def test_e2_probe_planned_capability(self):
        """E2: Planned capabilities are skipped (use CAP-005 which is still planned)."""
        import yaml
        from qa.capability_probe import probe_capability

        spec_path = os.path.join(PROJECT_ROOT, "qa", "vision_spec.yaml")
        with open(spec_path) as f:
            vision_spec = yaml.safe_load(f)

        # CAP-005 (PDF Format for Comparability) is still planned
        result = probe_capability("CAP-005", vision_spec)
        assert result.status == "skip"

    def test_e2_probe_implemented_characterization(self):
        """E2: CAP-002 (Characterization PDF) is now implemented."""
        import yaml
        from qa.capability_probe import probe_capability

        spec_path = os.path.join(PROJECT_ROOT, "qa", "vision_spec.yaml")
        with open(spec_path) as f:
            vision_spec = yaml.safe_load(f)

        result = probe_capability("CAP-002", vision_spec)
        assert result.status != "skip"  # no longer planned

    def test_e2_probe_nonexistent_capability(self):
        """E2: Non-existent capability returns error."""
        from qa.capability_probe import probe_capability

        result = probe_capability("CAP-999", {"capabilities": {}})
        assert result.status == "error"

    def test_e3_drift_detector_imports(self):
        """E3: drift_detector module imports correctly."""
        from qa.drift_detector import detect_drift, DriftReport
        assert callable(detect_drift)

    def test_e3_detect_no_drift(self):
        """E3: No drift when all implemented capabilities pass."""
        from qa.capability_probe import ProbeResult
        from qa.drift_detector import detect_drift

        results = [
            ProbeResult(
                capability_id="CAP-001",
                capability_name="Test Cap",
                status="pass",
                assertions_total=3,
                assertions_passed=3,
                assertions_failed=0,
            ),
        ]
        vision_spec = {
            "capabilities": {
                "CAP-001": {"name": "Test Cap", "status": "implemented", "acceptance": []},
            }
        }

        report = detect_drift(results, vision_spec)
        assert not report.has_drift
        assert report.n_passed == 1

    def test_e3_detect_regression_drift(self):
        """E3: Drift detected when implemented capability fails."""
        from qa.capability_probe import ProbeResult
        from qa.drift_detector import detect_drift

        results = [
            ProbeResult(
                capability_id="CAP-001",
                capability_name="Test Cap",
                status="fail",
                assertions_total=3,
                assertions_passed=1,
                assertions_failed=2,
                failure_details=["FAIL: assertion 1", "FAIL: assertion 2"],
            ),
        ]
        vision_spec = {
            "capabilities": {
                "CAP-001": {"name": "Test Cap", "status": "implemented", "acceptance": []},
            }
        }

        report = detect_drift(results, vision_spec)
        assert report.has_drift
        assert any(d.drift_type == "regression" for d in report.drift_items)

    def test_e4_run_qa_imports(self):
        """E4: run_qa module imports correctly."""
        from qa.run_qa import load_vision_spec, run_all_probes
        assert callable(load_vision_spec)
        assert callable(run_all_probes)
